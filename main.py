"""
City of Pacifica - Agenda Summary Generator
A modern desktop application for generating AI-powered council agenda summaries
"""

import customtkinter as ctk
import pandas as pd
from llama_cpp import Llama
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
import os
from datetime import datetime
import threading
from tkinter import ttk
import traceback

# Configure CustomTkinter
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# Required CSV headers
REQUIRED_HEADERS = [
    "MEETING DATE", "AGENDA ITEM", "DEPT", "AGENDA SECTION", "ACTION",
    "ATTY REVIEW REQ'D", "FIN REVIEW REQ'D", "DUE TO CAO / FIN", "DRAFT TO CM",
    "MAYOR MTG", "MT FINAL DUE", "AGENDA PUBLISH", "NOTES", "EXCEPTIONS",
    "NOTICING REQ;\nREGULATORY DEADLINES", "CONFLICT OF INTEREST REVIEW NEEDED (Y / N)",
    "HIGHLIGHT IN CWP", "Include in Summary for Mayor"
]

class AgendaSummaryGenerator(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # Window configuration
        self.title("City of Pacifica - Agenda Summary Generator")
        self.geometry("1280x720")
        self.minsize(800, 600)
        
        # Set window icon if available
        try:
            if os.path.exists("icon.ico"):
                self.iconbitmap("icon.ico")
        except:
            pass
        
        # Color theme
        self.bg_color = "#F5F5DC"
        self.primary_color = "#4682B4"
        self.accent_color = "#5A9BD5"
        self.text_color = "#333333"
        
        # Configure the main window background
        self.configure(fg_color=self.bg_color)
        
        # Data storage
        self.csv_data = None
        self.filtered_items = []
        self.selected_items = {}
        self.llm_model = None
        
        # Create main container
        self.main_container = ctk.CTkFrame(self, fg_color=self.bg_color)
        self.main_container.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Create navigation
        self.create_navigation()
        
        # Create views
        self.create_views()
        
        # Load LLM model in background
        self.load_llm_model()
        
    def create_navigation(self):
        """Create the top navigation bar"""
        nav_frame = ctk.CTkFrame(self.main_container, fg_color=self.bg_color, height=50)
        nav_frame.pack(fill="x", pady=(0, 20))
        nav_frame.pack_propagate(False)
        
        # Navigation buttons
        self.nav_buttons = {}
        button_names = ["Home", "Help", "Credits"]
        
        for i, name in enumerate(button_names):
            btn = ctk.CTkButton(
                nav_frame,
                text=name,
                width=120,
                height=40,
                fg_color=self.primary_color,
                hover_color=self.accent_color,
                text_color="white",
                command=lambda n=name: self.show_view(n)
            )
            btn.pack(side="left", padx=(0, 10))
            self.nav_buttons[name] = btn
        
    def create_views(self):
        """Create all view containers"""
        # Container for all views
        self.view_container = ctk.CTkFrame(self.main_container, fg_color=self.bg_color)
        self.view_container.pack(fill="both", expand=True)
        
        # Create individual views
        self.views = {}
        
        # Home view
        self.views["Home"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_home_view()
        
        # Help view
        self.views["Help"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_help_view()
        
        # Credits view
        self.views["Credits"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_credits_view()
        
        # Show Home view by default
        self.show_view("Home")
        
    def show_view(self, view_name):
        """Show the selected view and hide others"""
        for name, view in self.views.items():
            if name == view_name:
                view.pack(fill="both", expand=True)
            else:
                view.pack_forget()
                
    def create_home_view(self):
        """Create the Home view with Upload and Review states"""
        home_frame = self.views["Home"]
        
        # Upload state container
        self.upload_state = ctk.CTkFrame(home_frame, fg_color=self.bg_color)
        self.upload_state.pack(fill="both", expand=True)
        
        # Center container for upload elements
        center_frame = ctk.CTkFrame(self.upload_state, fg_color=self.bg_color)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        # City logo
        try:
            if os.path.exists("logo.png"):
                logo_image = Image.open("logo.png")
                logo_image = logo_image.resize((200, 200), Image.Resampling.LANCZOS)
                logo_ctk = ctk.CTkImage(logo_image, size=(200, 200))
                logo_label = ctk.CTkLabel(center_frame, image=logo_ctk, text="")
                logo_label.pack(pady=(0, 30))
        except:
            pass
            
        # Drag and drop area
        drop_frame = ctk.CTkFrame(
            center_frame,
            width=600,
            height=300,
            fg_color="white",
            border_width=2,
            border_color=self.primary_color
        )
        drop_frame.pack()
        drop_frame.pack_propagate(False)
        
        # Drop area content
        drop_label = ctk.CTkLabel(
            drop_frame,
            text="Drag and Drop your .csv file here",
            font=ctk.CTkFont(size=24, weight="bold"),
            text_color=self.text_color
        )
        drop_label.pack(pady=(80, 20))
        
        # Upload button
        upload_btn = ctk.CTkButton(
            drop_frame,
            text="Or Click to Select File",
            width=200,
            height=50,
            fg_color=self.primary_color,
            hover_color=self.accent_color,
            font=ctk.CTkFont(size=16),
            command=self.select_file
        )
        upload_btn.pack()
        
        # Format warning
        warning_text = ("Please ensure your .csv file is formatted correctly. It must include columns for "
                       "the meeting date, agenda item, notes, and a column named 'Include in Summary for Mayor' "
                       "to identify items for the report.")
        warning_label = ctk.CTkLabel(
            center_frame,
            text=warning_text,
            font=ctk.CTkFont(size=12),
            text_color=self.text_color,
            wraplength=600,
            justify="center"
        )
        warning_label.pack(pady=(20, 0))
        
        # Review state container (initially hidden)
        self.review_state = ctk.CTkFrame(home_frame, fg_color=self.bg_color)
        
    def create_review_state(self):
        """Create the review state UI"""
        # Clear existing content
        for widget in self.review_state.winfo_children():
            widget.destroy()
            
        # Top control bar
        control_bar = ctk.CTkFrame(self.review_state, fg_color=self.bg_color, height=60)
        control_bar.pack(fill="x", pady=(0, 20))
        control_bar.pack_propagate(False)
        
        # Item count label
        count_text = f"Review Items for Report ({len(self.filtered_items)} items found)"
        self.count_label = ctk.CTkLabel(
            control_bar,
            text=count_text,
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=self.text_color
        )
        self.count_label.pack(side="left", padx=(20, 0))
        
        # Generate button
        self.generate_btn = ctk.CTkButton(
            control_bar,
            text="Generate Report",
            width=150,
            height=40,
            fg_color=self.primary_color,
            hover_color=self.accent_color,
            font=ctk.CTkFont(size=16),
            command=self.generate_report
        )
        self.generate_btn.pack(side="right", padx=(0, 20))
        
        # Select/Deselect all buttons
        select_all_btn = ctk.CTkButton(
            control_bar,
            text="Select All",
            width=100,
            height=30,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            command=self.select_all_items
        )
        select_all_btn.pack(side="right", padx=5)
        
        deselect_all_btn = ctk.CTkButton(
            control_bar,
            text="Deselect All",
            width=100,
            height=30,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            command=self.deselect_all_items
        )
        deselect_all_btn.pack(side="right", padx=5)
        
        # Scrollable frame for items
        self.scroll_frame = ctk.CTkScrollableFrame(
            self.review_state,
            fg_color="white",
            corner_radius=10
        )
        self.scroll_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # Create item frames
        self.selected_items = {}
        for i, item in enumerate(self.filtered_items):
            self.create_item_frame(i, item)
            
    def create_item_frame(self, index, item):
        """Create a frame for each agenda item"""
        item_frame = ctk.CTkFrame(
            self.scroll_frame,
            fg_color="#f0f0f0",
            corner_radius=5,
            height=80
        )
        item_frame.pack(fill="x", padx=10, pady=5)
        item_frame.pack_propagate(False)
        
        # Checkbox
        var = tk.BooleanVar(value=True)
        self.selected_items[index] = var
        
        checkbox = ctk.CTkCheckBox(
            item_frame,
            text="",
            variable=var,
            width=20,
            checkbox_width=20,
            checkbox_height=20
        )
        checkbox.pack(side="left", padx=(10, 5), pady=10)
        
        # Content frame
        content_frame = ctk.CTkFrame(item_frame, fg_color="#f0f0f0")
        content_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        # Meeting date
        date_label = ctk.CTkLabel(
            content_frame,
            text=f"Meeting Date: {item['MEETING DATE']}",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color=self.text_color,
            anchor="w"
        )
        date_label.pack(anchor="w", pady=(10, 5))
        
        # Agenda item
        agenda_text = str(item['AGENDA ITEM'])[:150] + "..." if len(str(item['AGENDA ITEM'])) > 150 else str(item['AGENDA ITEM'])
        agenda_label = ctk.CTkLabel(
            content_frame,
            text=agenda_text,
            font=ctk.CTkFont(size=11),
            text_color=self.text_color,
            anchor="w",
            wraplength=800
        )
        agenda_label.pack(anchor="w", pady=(0, 10))
        
    def create_help_view(self):
        """Create the Help view"""
        help_frame = self.views["Help"]
        
        # Scrollable text area
        scroll_text = ctk.CTkScrollableFrame(help_frame, fg_color="white")
        scroll_text.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Help content
        help_content = """How to Use the Agenda Summary Generator

Step 1: Prepare Your CSV File
Ensure your source data is a .csv file with the correct headers. The application will check this for you. Your data must contain agenda items you want to summarize.

Step 2: Upload Your File
On the 'Home' tab, either drag your .csv file into the box or click the button to select it from your computer.

Step 3: Review and Select Items
After uploading, the app will display a list of all agenda items marked for inclusion. All items are selected by default. You can uncheck any item you wish to exclude from the final report.

Step 4: Generate the Report
Click the 'Generate Report' button. The application will use its built-in local AI to summarize each selected item. This may take a few moments depending on the number of items.

Step 5: Save Your Document
Once finished, a "Save As" window will appear. Choose a name and location for your new Microsoft Word (.docx) report. The document is now ready for your final review and any manual additions.

For a video demonstration, please visit: [Link to be added later]"""
        
        help_label = ctk.CTkLabel(
            scroll_text,
            text=help_content,
            font=ctk.CTkFont(size=14),
            text_color=self.text_color,
            justify="left",
            anchor="nw"
        )
        help_label.pack(anchor="nw", padx=20, pady=20)
        
    def create_credits_view(self):
        """Create the Credits view"""
        credits_frame = self.views["Credits"]
        
        # Center container
        center_frame = ctk.CTkFrame(credits_frame, fg_color=self.bg_color)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        # Title
        title_label = ctk.CTkLabel(
            center_frame,
            text="Agenda Summary Generator v1.0",
            font=ctk.CTkFont(size=28, weight="bold"),
            text_color=self.text_color
        )
        title_label.pack(pady=(0, 20))
        
        # Description
        desc_label = ctk.CTkLabel(
            center_frame,
            text="This tool was designed and developed to enhance organizational efficiency for the City of Pacifica.",
            font=ctk.CTkFont(size=14),
            text_color=self.text_color,
            wraplength=600
        )
        desc_label.pack(pady=(0, 30))
        
        # Credits
        credits_text = """Project Lead & Core Developer:
Nickolas Yang

Project Organization & Coordination:
Madeleine Hur

Acknowledgements:
This application utilizes the Gemma-2B-It language model from Google and the llama-cpp-python library.
All logos and trademarks are the property of their respective owners."""
        
        credits_label = ctk.CTkLabel(
            center_frame,
            text=credits_text,
            font=ctk.CTkFont(size=12),
            text_color=self.text_color,
            justify="center"
        )
        credits_label.pack()
        
    def select_file(self):
        """Handle file selection"""
        filename = filedialog.askopenfilename(
            title="Select CSV File",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if filename:
            self.process_csv_file(filename)
            
    def process_csv_file(self, filepath):
        """Process and validate the CSV file"""
        try:
            # Read CSV file
            self.csv_data = pd.read_csv(filepath)
            
            # Validate headers
            if not self.validate_csv_headers():
                messagebox.showerror(
                    "Invalid File Format",
                    "The selected CSV file is missing required columns. Please check the file against the template and try again."
                )
                return
                
            # Filter items marked for inclusion
            self.filtered_items = []
            for _, row in self.csv_data.iterrows():
                if str(row.get('Include in Summary for Mayor', '')).upper() == 'Y':
                    self.filtered_items.append(row)
                    
            if not self.filtered_items:
                messagebox.showwarning(
                    "No Items Found",
                    "No items were marked for inclusion in the Mayor's summary. Please check your CSV file."
                )
                return
                
            # Switch to review state
            self.upload_state.pack_forget()
            self.review_state.pack(fill="both", expand=True)
            self.create_review_state()
            
        except Exception as e:
            messagebox.showerror(
                "Error Reading File",
                f"An error occurred while reading the file: {str(e)}"
            )
            
    def validate_csv_headers(self):
        """Validate that all required headers are present"""
        if self.csv_data is None:
            return False
            
        csv_headers = list(self.csv_data.columns)
        for required_header in REQUIRED_HEADERS:
            if required_header not in csv_headers:
                return False
        return True
        
    def select_all_items(self):
        """Select all items in the review list"""
        for var in self.selected_items.values():
            var.set(True)
            
    def deselect_all_items(self):
        """Deselect all items in the review list"""
        for var in self.selected_items.values():
            var.set(False)
            
    def load_llm_model(self):
        """Load the LLM model in background"""
        def load_model():
            try:
                model_path = "gemma-2b-it.gguf"
                if os.path.exists(model_path):
                    self.llm_model = Llama(
                        model_path=model_path,
                        n_ctx=2048,
                        n_threads=4,
                        n_gpu_layers=0
                    )
            except Exception as e:
                print(f"Failed to load LLM model: {e}")
                
        threading.Thread(target=load_model, daemon=True).start()
        
    def generate_report(self):
        """Generate the report with LLM summaries"""
        # Get selected items
        selected_data = []
        for idx, var in self.selected_items.items():
            if var.get():
                selected_data.append(self.filtered_items[idx])
                
        if not selected_data:
            messagebox.showwarning(
                "No Items Selected",
                "Please select at least one item to generate a report."
            )
            return
            
        # Disable generate button
        self.generate_btn.configure(text="Generating...", state="disabled")
        
        # Generate report in background
        threading.Thread(
            target=self._generate_report_thread,
            args=(selected_data,),
            daemon=True
        ).start()
        
    def _generate_report_thread(self, selected_data):
        """Background thread for report generation"""
        try:
            # Group items by meeting date and section
            grouped_items = {}
            for item in selected_data:
                date = item['MEETING DATE']
                section = item.get('AGENDA SECTION', 'Unknown Section')
                
                if date not in grouped_items:
                    grouped_items[date] = {}
                if section not in grouped_items[date]:
                    grouped_items[date][section] = []
                    
                grouped_items[date][section].append(item)
                
            # Generate summaries
            all_summaries = []
            meeting_dates = []
            
            for date in sorted(grouped_items.keys()):
                meeting_dates.append(date)
                all_summaries.append(f"Meeting Date: {date}")
                
                for section in sorted(grouped_items[date].keys()):
                    all_summaries.append(f"\n{section}:")
                    
                    for item in grouped_items[date][section]:
                        summary = self.generate_llm_summary(item)
                        all_summaries.append(f"• {summary}")
                        
                all_summaries.append("")  # Add spacing between dates
                
            # Generate document
            content = "\n".join(all_summaries)
            doc = self.create_word_document(content, meeting_dates)
            
            # Save document (switch to main thread)
            self.after(0, self._save_document, doc)
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror(
                "Generation Error",
                f"An error occurred while generating the report: {str(e)}"
            ))
            
        finally:
            # Re-enable button
            self.after(0, lambda: self.generate_btn.configure(
                text="Generate Report",
                state="normal"
            ))
            
    def generate_llm_summary(self, item):
        """Generate summary for a single agenda item"""
        if self.llm_model is None:
            # Fallback if model not loaded
            return f"{item['AGENDA ITEM']} - {item.get('NOTES', 'No notes available')}"
            
        try:
            # Prepare prompt
            prompt = f"""You are an expert municipal analyst. Summarize this agenda item in 2-3 concise sentences:

AGENDA ITEM: {item['AGENDA ITEM']}
DEPARTMENT: {item.get('DEPT', 'N/A')}
NOTES: {item.get('NOTES', 'No notes available')}

Summary:"""
            
            # Generate response
            response = self.llm_model(
                prompt,
                max_tokens=150,
                temperature=0.3,
                stop=["\\n\\n"]
            )
            
            return response['choices'][0]['text'].strip()
            
        except:
            # Fallback to simple concatenation
            return f"{item['AGENDA ITEM']} - {item.get('NOTES', 'No notes available')}"
            
    def create_word_document(self, content, meeting_dates):
        """Create the Word document with proper formatting"""
        doc = Document()
        
        # Add content
        doc.add_paragraph()  # Blank line
        
        # Updated date
        current_date = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(f"Updated {current_date}")
        
        # Calculate month range
        if meeting_dates:
            dates = [datetime.strptime(d, "%m/%d/%Y") if "/" in d else datetime.now() for d in meeting_dates]
            min_date = min(dates)
            max_date = max(dates)
            
            if min_date.month == max_date.month:
                month_range = min_date.strftime("%B %Y")
            else:
                month_range = f"{min_date.strftime('%B')} - {max_date.strftime('%B %Y')}"
        else:
            month_range = datetime.now().strftime("%B %Y")
            
        # Title
        title = doc.add_paragraph(f"Major Council Agenda Items, Tentative for {month_range}")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(16)
        
        # Note
        note_text = ("Note: This is a Tentative Agenda Listing. Dates of items are subject to change "
                    "up to the last minute for a variety of reasons. In addition, this listing does not "
                    "necessarily report all items, just ones that are noteworthy. The City Manager typically "
                    "reviews the tentative agenda items list in more detail with each Councilmember during "
                    "individual meetings.")
        note = doc.add_paragraph(note_text)
        note.runs[0].italic = True
        
        # Add horizontal line
        doc.add_paragraph("_" * 80)
        
        # Add LLM content
        doc.add_paragraph(content)
        
        # Add horizontal line
        doc.add_paragraph("_" * 80)
        
        # TBD section
        tbd = doc.add_paragraph("TBD:")
        tbd.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        doc.add_paragraph()
        
        # Significant items section
        last_report_date = (datetime.now() - pd.Timedelta(days=30)).strftime("%B %d, %Y")
        sig_items = doc.add_paragraph(f"Significant Items Completed Since {last_report_date}:")
        sig_items.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        
        return doc
        
    def _save_document(self, doc):
        """Save the document with file dialog"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            initialfile=f"Council_Agenda_Summary_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        
        if filename:
            try:
                doc.save(filename)
                messagebox.showinfo(
                    "Success",
                    f"Report generated successfully and saved to:\n{filename}"
                )
            except Exception as e:
                messagebox.showerror(
                    "Save Error",
                    f"Failed to save document: {str(e)}"
                )


def main():
    """Main entry point"""
    app = AgendaSummaryGenerator()
    app.mainloop()


if __name__ == "__main__":
    main()