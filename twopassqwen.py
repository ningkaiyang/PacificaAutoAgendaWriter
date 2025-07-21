"""
City of Pacifica - Agenda Summary Generator
A modern desktop application for generating AI-powered council agenda summaries
TODO: More preprocessing to make LLM performance better, Easier Format, fix bugs like scrolling and drag and drop and Save Report button taking from text box and Qwen Thinking Traces appearing 
"""

import customtkinter as ctk
import pandas as pd
from llama_cpp import Llama
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from PIL import Image
import os
from datetime import datetime
import threading
from tkinter import ttk
import traceback
import sys
import time
import contextlib
import resource  # i'll add this to track memory usage

# Configure CustomTkinter
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# Helper utilities from llama_demo.py
def logical_cores() -> int:
    """Return logical core count; always ≥ 1."""
    return max(1, os.cpu_count() or 1)

def default_threads() -> int:
    """Half the logical cores, minimum 1."""
    return max(1, logical_cores() // 2)

@contextlib.contextmanager
def suppress_stderr():
    """Temporarily suppress stderr output."""
    with open(os.devnull, "w") as devnull:
        old_stderr = sys.stderr
        sys.stderr = devnull
        try:
            yield
        finally:
            sys.stderr = old_stderr

class TokenStreamer:
    """Collects streamed tokens, prints all tokens (including thinking tags) for debugging,
    and tracks speed."""
    def __init__(self):
        self._start = time.perf_counter()
        self._tok = 0
        
    def __call__(self, chunk: dict):
        tok = chunk["choices"][0]["delta"].get("content", "")
        if not tok:
            return
        self._tok += 1
        # Print everything to console for debugging
        print(tok, end="", flush=True)
        
    def done(self):
        dt = time.perf_counter() - self._start
        if dt:
            print(f"\nAverage speed: {self._tok/dt:.2f} tok/s")
            print(f"Tokens: {self._tok}")
            print(f"Elapsed Time: {dt:.2f}s")
            # memory usage on macOS is in bytes, so i'll convert to MB
            mem_mb = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss / 1024 / 1024
            print(f"Peak Memory Usage: {mem_mb:.2f} MB")

class GUITokenFilter:
    """Filters out <think>...</think> blocks for clean GUI display."""
    def __init__(self):
        self._buf = ""
        self._in_think = False
        
    def filter_token(self, token: str) -> str:
        """Filter a token and return the clean portion for GUI display."""
        if not token:
            return ""
            
        self._buf += token
        result = ""
        
        while self._buf:
            if not self._in_think:
                # Look for start of thinking block
                think_start = self._buf.find("<think>")
                if think_start == -1:
                    # No thinking block found, return all buffered content
                    result += self._buf
                    self._buf = ""
                else:
                    # Found thinking block start, return content before it
                    result += self._buf[:think_start]
                    self._buf = self._buf[think_start:]
                    self._in_think = True
                    # Remove the <think> tag from buffer
                    if self._buf.startswith("<think>"):
                        self._buf = self._buf[7:]  # len("<think>") = 7
            else:
                # We're inside a thinking block, look for end
                think_end = self._buf.find("</think>")
                if think_end == -1:
                    # No end found yet, consume all buffer
                    self._buf = ""
                else:
                    # Found end, remove everything up to and including </think>
                    self._buf = self._buf[think_end + 8:]  # len("</think>") = 8
                    self._in_think = False
                    
        return result

PROMPT_TEMPLATE = """You are an expert city clerk responsible for creating agenda summaries for the City Council. Your task is to take a list of agenda items for a specific meeting date and format them into a clear, concise report.

You have recieved a set of summarized items. You are to categorize them and properly edit them in small ways if necessary (capitalization, merging, etc) and put them together into a resport.
Follow these rules strictly:
1.  Format: The output must be raw text only. Do not use any markdown like '##' or '**'.
2.  Date Header: The report must start with the FULL month name followed by the day number, e.g. "January 1:".  NEVER use numeric-month abbreviations such as "1-Jan".  If there are meeting-level notes, place them in parentheses immediately after the date.
3.  Sections: The report must BEGIN with EITHER "Study Session:" or "Closed Session:" depending on which type of item exists for that meeting date.
        • If BOTH exist, list "Study Session:" first and "Closed Session:" second.
        • If neither exists, omit them and start with the first section that does have items.
    After the opening section(s) continue with these sections which MUST be included, and in the following order:
        "Special Presentations:"
        "Consent:"
        "Consideration or Public Hearing:"
    If a section has no items, write "TBD" right after the section name. Example: "Closed Session: TBD" or "Consent: TBD"
4.  Item Bullet Points:
    - CRITICAL: Each individual agenda item provided to you MUST be on its own new line in the output.
    - Every item's line must start with a single hyphen and a space: "- ". Do NOT use other bullet point characters like '•' to start off a new line.

Here are some examples of the desired output format:

Example 1:
June 23:
Closed Session: TBD
Special Presentations:
- Parks Make Life Better Month
Consent:
- Childcare site lease agreement with Pacifica School District
- New Operating Agreement with PRC for TSPP
- Design Services Agreement for FY26-27 Pavement Resurfacing Project
- Recertification of Sewer System Management Plan (State law requirement)
Consideration or Public Hearing:
- FY 2025-26 Budget Adoption
- Annual position vacancy, recruitment and retention report (State law requirement AB2561)
- Introduction of Ordinance Changing Council Meeting start-time and formal adoption of other Governance Training outcomes

Example 2:
July 14:
Closed Session: TBD
Special Presentations:
- Joann Arnos, OSPAC Years of Service
Consent:
- Annual POs/Agreements over $75K PWD-Wastewater
- Labor MOUs (placeholder)
- Sewer service charges for FY2025-26 (last year of approved 5-year schedule)
- VRBO Voluntary Collection Agreement (placeholder)
Consideration or Public Hearing:
- STR Ordinance Update Introduction
- Continued Consideration of Climate Action and Resilience Plan Adoption

Example 3 (Handling pending descriptions and TBDs):
August 5:
Closed Session: TBD
Special Presentations: TBD
Consent:
- Resolution for park naming - ADD DESCRIPTION
Consideration or Public Hearing: TBD

Example 4 (Meeting that includes both a Closed Session and Study Session):
September 10:
Study Session:
- Joint Study Session on Revenue Options - ADD DESCRIPTION
Closed Session: TBD
Special Presentations: TBD
Consent:
- Bi-Weekly Disbursements approval
Consideration or Public Hearing: TBD

Now, generate a report for the following meeting date based on the items provided below. Remember to place each item on a new line and to summarize each item to a few sentences.

Meeting Date: {meeting_date} - IMPORTANT: THIS IS THE ACTUAL METING DATE FOR YOUR REPORT!!!

Agenda Items (pre-sorted by section):
{items_text}

Report:
"""

# Required CSV headers
REQUIRED_HEADERS = [
    "MEETING DATE", "AGENDA ITEM", "DEPT", "AGENDA SECTION", "ACTION",
    "ATTY REVIEW REQ'D", "FIN REVIEW REQ'D", "DUE TO CAO / FIN", "DRAFT TO CM",
    "MAYOR MTG", "MT FINAL DUE", "AGENDA PUBLISH", "NOTES", "EXCEPTIONS",
    "NOTICING REQ;\nREGULATORY DEADLINES", "CONFLICT OF INTEREST REVIEW NEEDED (Y / N)",
    "HIGHLIGHT IN CWP", "Include in Summary for Mayor"
]

class AgendaSummaryGenerator(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # Window configuration
        self.title("City of Pacifica - Agenda Summary Generator")
        self.geometry("1280x720")
        self.minsize(800, 600)
        
        # Set window icon if available
        try:
            if os.path.exists("icon.ico"):
                self.iconbitmap("icon.ico")
        except:
            pass
        
        # Color theme
        self.bg_color = "#F5F5DC"
        self.primary_color = "#4682B4"
        self.accent_color = "#5A9BD5"
        self.text_color = "#333333"
        
        # Configure the main window background
        self.configure(fg_color=self.bg_color)
        
        # Data storage
        self.csv_data = None
        self.filtered_items = []
        self.tree = None # Replaces self.selected_items
        self.llm_model = None
        self.generated_report_text = ""
        self.meeting_dates_for_report = []
        
        # Create main container
        self.main_container = ctk.CTkFrame(self, fg_color=self.bg_color)
        self.main_container.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Create navigation
        self.create_navigation()
        
        # Create views
        self.create_views()
        
        # Load LLM model in background
        self.load_llm_model()
        
    def create_navigation(self):
        """Create the top navigation bar"""
        nav_frame = ctk.CTkFrame(self.main_container, fg_color=self.bg_color, height=50)
        nav_frame.pack(fill="x", pady=(0, 20))
        nav_frame.pack_propagate(False)
        
        # Navigation buttons
        self.nav_buttons = {}
        button_names = ["Home", "Help", "Credits"]
        
        for i, name in enumerate(button_names):
            btn = ctk.CTkButton(
                nav_frame,
                text=name,
                width=120,
                height=40,
                fg_color=self.primary_color,
                hover_color=self.accent_color,
                text_color="white",
                command=lambda n=name: self.show_view(n)
            )
            btn.pack(side="left", padx=(0, 10))
            self.nav_buttons[name] = btn
        
    def create_views(self):
        """Create all view containers"""
        # Container for all views
        self.view_container = ctk.CTkFrame(self.main_container, fg_color=self.bg_color)
        self.view_container.pack(fill="both", expand=True)
        
        # Create individual views
        self.views = {}
        
        # Home view
        self.views["Home"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_home_view()
        
        # Help view
        self.views["Help"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_help_view()
        
        # Credits view
        self.views["Credits"] = ctk.CTkFrame(self.view_container, fg_color=self.bg_color)
        self.create_credits_view()
        
        # Show Home view by default
        self.show_view("Home")
        
    def show_view(self, view_name):
        """Show the selected view and hide others"""
        # If going "Home" from nav, always reset to the initial upload screen.
        if view_name == "Home":
            for name, view in self.views.items():
                if name != "Home":
                    view.pack_forget()
            self.views["Home"].pack(fill="both", expand=True)
            # Ensure the upload state is shown and others are hidden within Home
            self.review_state.pack_forget()
            self.generation_state.pack_forget()
            self.upload_state.pack(fill="both", expand=True)
        else:
            for name, view in self.views.items():
                if name == view_name:
                    view.pack(fill="both", expand=True)
                else:
                    view.pack_forget()
                
    def create_home_view(self):
        """Create the Home view with Upload and Review states"""
        home_frame = self.views["Home"]
        
        # Upload state container
        self.upload_state = ctk.CTkFrame(home_frame, fg_color=self.bg_color)
        self.upload_state.pack(fill="both", expand=True)
        
        # Center container for upload elements
        center_frame = ctk.CTkFrame(self.upload_state, fg_color=self.bg_color)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        # City logo
        try:
            if os.path.exists("logo.png"):
                logo_image = Image.open("logo.png")
                logo_image = logo_image.resize((200, 200), Image.Resampling.LANCZOS)
                logo_ctk = ctk.CTkImage(logo_image, size=(200, 200))
                logo_label = ctk.CTkLabel(center_frame, image=logo_ctk, text="")
                logo_label.pack(pady=(0, 30))
        except:
            pass
            
        # Drag and drop area
        drop_frame = ctk.CTkFrame(
            center_frame,
            width=600,
            height=300,
            fg_color="white",
            border_width=2,
            border_color=self.primary_color
        )
        drop_frame.pack()
        drop_frame.pack_propagate(False)
        
        # Drop area content
        drop_label = ctk.CTkLabel(
            drop_frame,
            text="Drag & Drop not fully supported. Please use the button.",
            font=ctk.CTkFont(size=20, weight="bold"),
            text_color=self.text_color,
            wraplength=550
        )
        drop_label.pack(pady=(80, 20))
        
        # Upload button
        upload_btn = ctk.CTkButton(
            drop_frame,
            text="Or Click to Select File",
            width=200,
            height=50,
            fg_color=self.primary_color,
            hover_color=self.accent_color,
            font=ctk.CTkFont(size=16),
            command=self.select_file
        )
        upload_btn.pack()
        
        # Format warning
        warning_text = ("Please ensure your .csv file is formatted correctly. It must include columns for "
                       "the meeting date, agenda item, notes, and a column named 'Include in Summary for Mayor' "
                       "to identify items for the report.")
        warning_label = ctk.CTkLabel(
            center_frame,
            text=warning_text,
            font=ctk.CTkFont(size=12),
            text_color=self.text_color,
            wraplength=600,
            justify="center"
        )
        warning_label.pack(pady=(20, 0))
        
        # Review state container (initially hidden)
        self.review_state = ctk.CTkFrame(home_frame, fg_color=self.bg_color)
        
        # Generation state container (initially hidden)
        self.generation_state = ctk.CTkFrame(home_frame, fg_color=self.bg_color)
        self.create_generation_state_widgets()
        
    def create_generation_state_widgets(self):
        """Create the UI for the report generation state"""
        # Top control bar
        control_bar = ctk.CTkFrame(self.generation_state, fg_color=self.bg_color, height=60)
        control_bar.pack(fill="x", pady=(0, 20))
        control_bar.pack_propagate(False)

        # Add a back button to go to review screen
        back_btn = ctk.CTkButton(
            control_bar,
            text="Back to Review",
            width=150,
            height=40,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            font=ctk.CTkFont(size=16),
            command=self.back_to_review_state
        )
        back_btn.pack(side="left", padx=20)
        
        # Title
        title_label = ctk.CTkLabel(
            control_bar,
            text="Generating Report...",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=self.text_color
        )
        title_label.pack(side="left", padx=20)
        
        # Save button (initially disabled)
        self.save_btn = ctk.CTkButton(
            control_bar,
            text="Save Report",
            width=150,
            height=40,
            fg_color=self.primary_color,
            hover_color=self.accent_color,
            font=ctk.CTkFont(size=16),
            command=self.save_generated_report,
            state="disabled"
        )
        self.save_btn.pack(side="right", padx=20)
        
        # Text box for streaming output
        self.generation_textbox = ctk.CTkTextbox(
            self.generation_state,
            fg_color="white",
            text_color=self.text_color,
            font=ctk.CTkFont(family="monospace", size=12),
            wrap="word"
        )
        self.generation_textbox.pack(fill="both", expand=True, padx=20, pady=(0, 20))

    def back_to_review_state(self):
        """Go from generation screen back to review screen."""
        self.generation_state.pack_forget()
        self.review_state.pack(fill="both", expand=True)

    def create_review_state(self):
        """Create the review state UI"""
        # Clear existing content
        for widget in self.review_state.winfo_children():
            widget.destroy()
            
        # Top control bar
        control_bar = ctk.CTkFrame(self.review_state, fg_color=self.bg_color, height=60)
        control_bar.pack(fill="x", pady=(0, 20))
        control_bar.pack_propagate(False)

        # Add Back button
        back_btn = ctk.CTkButton(
            control_bar,
            text="Back to Home",
            width=120,
            height=30,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            command=lambda: self.show_view("Home")
        )
        back_btn.pack(side="left", padx=10)
        
        # Item count label
        count_text = f"Review Items for Report ({len(self.filtered_items)} items found)"
        self.count_label = ctk.CTkLabel(
            control_bar,
            text=count_text,
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=self.text_color
        )
        self.count_label.pack(side="left", padx=(20, 0))
        
        # Generate button
        self.generate_btn = ctk.CTkButton(
            control_bar,
            text="Generate Report",
            width=150,
            height=40,
            fg_color=self.primary_color,
            hover_color=self.accent_color,
            font=ctk.CTkFont(size=16),
            command=self.generate_report
        )
        self.generate_btn.pack(side="right", padx=(0, 20))
        
        # Select/Deselect all buttons
        select_all_btn = ctk.CTkButton(
            control_bar,
            text="Select All",
            width=100,
            height=30,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            command=self.select_all_items
        )
        select_all_btn.pack(side="right", padx=5)
        
        deselect_all_btn = ctk.CTkButton(
            control_bar,
            text="Deselect All",
            width=100,
            height=30,
            fg_color=self.accent_color,
            hover_color=self.primary_color,
            command=self.deselect_all_items
        )
        deselect_all_btn.pack(side="right", padx=5)

        # --- START: Treeview Implementation ---
        # This replaces the slow, chunky CTkScrollableFrame of individual widgets
        style = ttk.Style(self)
        style.theme_use("default")
        # Configure Treeview colors and row height for a compact look
        style.configure("Treeview",
                        background="white",
                        foreground=self.text_color,
                        fieldbackground="white",
                        rowheight=25,
                        font=ctk.CTkFont(size=11))
        style.map('Treeview',
                  background=[('selected', self.primary_color)],
                  foreground=[('selected', 'white')])
        style.configure("Treeview.Heading", font=ctk.CTkFont(size=12, weight="bold"))

        tree_container = ctk.CTkFrame(self.review_state, fg_color="transparent")
        tree_container.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        self.tree = ttk.Treeview(
            tree_container,
            columns=("Section", "Item", "Notes"),
            show="headings",
            selectmode="extended" # Allows multiple selections
        )

        # Define headings
        self.tree.heading("Section", text="Section", anchor='w')
        self.tree.heading("Item", text="Agenda Item", anchor='w')
        self.tree.heading("Notes", text="Notes", anchor='w')

        # Define column properties
        self.tree.column("Section", width=180, stretch=False, anchor='w')
        self.tree.column("Item", width=500, stretch=True, anchor='w')
        self.tree.column("Notes", width=350, stretch=True, anchor='w')

        # Populate treeview with data
        for i, item in enumerate(self.filtered_items):
            # Clean up data for display
            section = str(item.get('AGENDA SECTION', 'N/A')).replace('\n', ' ')
            agenda_item = str(item.get('AGENDA ITEM', 'N/A')).replace('\n', ' ')
            notes = str(item.get('NOTES', '')).replace('\n', ' ')
            # Use index 'i' as the unique item identifier (iid)
            self.tree.insert("", "end", values=(section, agenda_item, notes), iid=str(i))

        # Add a scrollbar that works with mouse/trackpad
        scrollbar = ttk.Scrollbar(tree_container, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        # Pack everything
        scrollbar.pack(side="right", fill="y")
        self.tree.pack(side="left", fill="both", expand=True)

        # Select all items by default on first load
        self.select_all_items()

        # Bind click to toggle selection for better UX
        self.tree.bind("<Button-1>", self.on_tree_click)
        # --- END: Treeview Implementation ---
            
    def on_tree_click(self, event):
        """Handle clicks on the treeview to select/deselect rows."""
        region = self.tree.identify("region", event.x, event.y)
        # Only toggle selection if a cell (not heading) is clicked
        if region == "cell":
            row_id = self.tree.identify_row(event.y)
            if self.tree.exists(row_id):
                if row_id in self.tree.selection():
                    self.tree.selection_remove(row_id)
                else:
                    self.tree.selection_add(row_id)
            
    def create_help_view(self):
        """Create the Help view"""
        help_frame = self.views["Help"]
        
        # Scrollable text area
        scroll_text = ctk.CTkScrollableFrame(help_frame, fg_color="white")
        scroll_text.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Help content
        help_content = """How to Use the Agenda Summary Generator

Step 1: Prepare Your CSV File
Ensure your source data is a .csv file with the correct headers. The application will check this for you. Your data must contain agenda items you want to summarize.

Step 2: Upload Your File
On the 'Home' tab, either drag your .csv file into the box or click the button to select it from your computer.

Step 3: Review and Select Items
After uploading, the app will display a list of all agenda items marked for inclusion. All items are selected by default. You can uncheck any item you wish to exclude from the final report.

Step 4: Generate the Report
Click the 'Generate Report' button. The application will use its built-in local AI to summarize each selected item. This may take a few moments depending on the number of items.

Step 5: Save Your Document
Once finished, a "Save As" window will appear. Choose a name and location for your new Microsoft Word (.docx) report. The document is now ready for your final review and any manual additions.

For a video demonstration, please visit: [Link to be added later]"""
        
        help_label = ctk.CTkLabel(
            scroll_text,
            text=help_content,
            font=ctk.CTkFont(size=14),
            text_color=self.text_color,
            justify="left",
            anchor="nw"
        )
        help_label.pack(anchor="nw", padx=20, pady=20)
        
    def create_credits_view(self):
        """Create the Credits view"""
        credits_frame = self.views["Credits"]
        
        # Center container
        center_frame = ctk.CTkFrame(credits_frame, fg_color=self.bg_color)
        center_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        # Title
        title_label = ctk.CTkLabel(
            center_frame,
            text="Agenda Summary Generator v1.0",
            font=ctk.CTkFont(size=28, weight="bold"),
            text_color=self.text_color
        )
        title_label.pack(pady=(0, 20))
        
        # Description
        desc_label = ctk.CTkLabel(
            center_frame,
            text="This tool was designed and developed to enhance organizational efficiency for the City of Pacifica.",
            font=ctk.CTkFont(size=14),
            text_color=self.text_color,
            wraplength=600
        )
        desc_label.pack(pady=(0, 30))
        
        # Credits
        credits_text = """Project Lead & Core Developer:
Nickolas Yang

Project Organization & Coordination:
Madeleine Hur

Acknowledgements:
This application utilizes open-source Large Language Models along with the llama-cpp-python library.
All logos and trademarks are the property of their respective owners.
Chatbot icon created by juicy_fish - Flaticon."""
        
        credits_label = ctk.CTkLabel(
            center_frame,
            text=credits_text,
            font=ctk.CTkFont(size=12),
            text_color=self.text_color,
            justify="center"
        )
        credits_label.pack()
        
    def select_file(self):
        """Handle file selection"""
        filename = filedialog.askopenfilename(
            title="Select CSV File",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if filename:
            self.process_csv_file(filename)
            
    def process_csv_file(self, filepath):
        """Process and validate the CSV file"""
        try:
            # Read CSV file
            self.csv_data = pd.read_csv(filepath)
            
            # Validate headers
            if not self.validate_csv_headers():
                messagebox.showerror(
                    "Invalid File Format",
                    "The selected CSV file is missing required columns. Please check the file against the template and try again."
                )
                return
                
            # Filter items marked for inclusion
            self.filtered_items = []
            for _, row in self.csv_data.iterrows():
                if str(row.get('Include in Summary for Mayor', '')).upper() == 'Y':
                    self.filtered_items.append(row)
                    
            if not self.filtered_items:
                messagebox.showwarning(
                    "No Items Found",
                    "No items were marked for inclusion in the Mayor's summary. Please check your CSV file."
                )
                return
                
            # Switch to review state
            self.upload_state.pack_forget()
            self.review_state.pack(fill="both", expand=True)
            self.create_review_state()
            
        except Exception as e:
            messagebox.showerror(
                "Error Reading File",
                f"An error occurred while reading the file: {str(e)}"
            )
            
    def validate_csv_headers(self):
        """Validate that all required headers are present"""
        if self.csv_data is None:
            return False
            
        csv_headers = list(self.csv_data.columns)
        for required_header in REQUIRED_HEADERS:
            if required_header not in csv_headers:
                return False
        return True
        
    def select_all_items(self):
        """Select all items in the review list"""
        if self.tree:
            children = self.tree.get_children()
            self.tree.selection_set(children)
            
    def deselect_all_items(self):
        """Deselect all items in the review list"""
        if self.tree:
            self.tree.selection_set()
            
    def load_llm_model(self):
        """Load the LLM model in background"""
        def load_model():
            try:
                model_path = "language_models/Qwen3-4B-Q6_K.gguf"
                if os.path.exists(model_path):
                    with suppress_stderr():
                        self.llm_model = Llama(
                            model_path=model_path,
                            chat_format="chatml",
                            n_ctx=10000,
                            n_threads=default_threads(),
                            verbose=False,
                        )
            except Exception as e:
                print(f"Failed to load LLM model: {e}")
                
        threading.Thread(target=load_model, daemon=True).start()
        
    def generate_report(self):
        """Generate the report with LLM summaries"""
        # Get selected items from the Treeview
        if not self.tree:
            messagebox.showwarning("No Items Selected", "Please select at least one item to generate a report.")
            return

        selected_iids = self.tree.selection()
        if not selected_iids:
            messagebox.showwarning("No Items Selected", "Please select at least one item to generate a report.")
            return

        # Map the selected iids (which are string indices) back to the original data
        selected_data = [self.filtered_items[int(iid)] for iid in selected_iids]
            
        # Switch to generation view
        self.review_state.pack_forget()
        self.generation_state.pack(fill="both", expand=True)
        self.generation_textbox.delete("1.0", "end")
        self.save_btn.configure(state="disabled")
        
        # Disable generate button
        self.generate_btn.configure(text="Generating...", state="disabled")
        
        # Generate report in background
        threading.Thread(
            target=self._generate_report_thread,
            args=(selected_data,),
            daemon=True
        ).start()
        
    def _generate_report_thread(self, selected_data):
        """Background thread for report generation using streaming."""
        try:
            self.generated_report_text = ""
            
            # Group items by meeting date while preserving the order from the CSV
            grouped_by_date = {}
            meeting_dates_in_order = []
            for item in selected_data:
                date = item['MEETING DATE']
                if date not in grouped_by_date:
                    grouped_by_date[date] = []
                    meeting_dates_in_order.append(date)
                grouped_by_date[date].append(item)
            
            self.meeting_dates_for_report = meeting_dates_in_order

            if self.llm_model is None:
                raise ConnectionError("LLM model is not loaded. Please wait or restart the application.")

            # Process each date sequentially
            for date in self.meeting_dates_for_report:
                items = grouped_by_date[date]

                # Pre-sort items by agenda section to help the LLM
                items.sort(key=lambda x: str(x.get('AGENDA SECTION', '')))
                
                items_text = ""
                for item in items:
                    section = str(item.get('AGENDA SECTION', 'N/A')).replace('\n', ' ')
                    agenda_item = str(item.get('AGENDA ITEM', 'N/A')).replace('\n', ' ')
                    notes = str(item.get('NOTES', 'No notes available')).replace('\n', ' ')
                    items_text += f"- Section: {section}, Item: \"{agenda_item}\", Notes: \"{notes}\"\n"
                
                # --- START: TWO-PASS GENERATION ---

                # PASS 1: Generate line-by-line summaries
                summarization_prompt = f"""You are an expert city clerk. Your task is to summarize each agenda item into ONE short clause (around 15 words or fewer). Don't waste time counting words too hard though.

Rules for summarization:
- Summarize each agenda item in ONE short clause that clearly signals what the item is
- You MUST omit unnecessary internal workflow words such as "moved from [dates]", and "per [person]". DO NOT say "moved from 1/1 to 12/31 per Y.Carter" or "per K.Woodhouse" or such.
- Remove characters that would not work well for reading within the item like all "•" characters
- If an item has multiple details, combine them using parentheses "()" or semicolons ";" ONLY, do not use bullets "•"
- If an item includes "placeholder", append "(placeholder)" with no other unnecessary placeholder details to the end
- If an item include "ADD DESCRIPTION", delete it and append " - ADD DESCRIPTION" to the end, after any potential "placeholder"
- The summaries should be prepended by which category they belong in: "Study Session:" or "Closed Session:" or "Special Presentations:" or "Consent:" or "Consideration or Public Hearing:".
- Each summary title must use Title Case (capitalize all principal words), for example: "Approval of Minutes for 1/1/2025 Meeting"
- If an agenda item includes a date range in mm/dd/yyyy format, preserve that exact format for conciseness; only the meeting date header should be written out in full word form.

Some good examples:
- Special Presentations: Proclamation - Suicide Prevention Month - September 2025 - ADD DESCRIPTION
- Closed Study: TBD
- Study Session: Study Session on Revenue Generation - ADD DESCRIPTION
- Special Presentations: City Staff New Hires (Semi-Annual Update) - moved from 6/23 to 8/25 per Y.Carter
- Consent: Annual POs/Agreements over $75K PWD-Wastewater
- Consent: Sewer service charges for FY2025-26 (last year of approved 5-year schedule)
- Consideration or Public Hearing: Continued Consideration of Climate Action and Resilience Plan Adoption

Meeting Date: {date} - IMPORTANT! THIS IS THE ACTUAL MEETING DATE, KEEP TRACK OF IT CAREFULLY! Start off your summarization lines with this meeting date, parsed neatly as <Month Day> like "Meeting Date: January 1" or "Meeting Date: December 31".

Agenda Items to Summarize:
{items_text.strip()}

IMPORTANT: Note you only have a 1000 token limit before you must create an output, so don't think in circles and just generate summaries and output when they look decent. Save tokens and decide on a summary per line quickly without overthinking.
Provide ONLY the proper meeting date format and then the summarized lines, each CAREFULLY capitalized and prepended properly, one per line: /think"""
                
                print("\n--- PASS 1: SUMMARIZATION ---")
                summarization_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": summarization_prompt}],
                    max_tokens=4000,  # limited tokens for focused summarization
                    temperature=0,  # lower temperature for more consistent summaries
                    top_p=0.95,
                    top_k=20,
                    stream=True,
                )

                # Collect summarized items from Pass 1
                think_streamer = TokenStreamer()
                summarized_items = ""
                for chunk in summarization_stream:
                    token = chunk["choices"][0]["delta"].get("content", "")
                    summarized_items += token
                    think_streamer(chunk)  # count tokens and print for debug
                think_streamer.done()

                # Clean up summarized_items to remove any incomplete thinking tags
                # and extract only the actual bullet point content
                clean_summarized_items = self._extract_clean_summary(summarized_items)

                # PASS 2: Generate final formatted report using the summaries
                # use the main prompt template for the final formatting pass
                # we replace the original items_text with the summarized ones from pass 1
                # and add a /no_think instruction to prevent the model from re-summarizing
                format_prompt = PROMPT_TEMPLATE.format(
                    meeting_date=date,
                    items_text=clean_summarized_items.strip()
                ) + ' /no_think'

                print("\n--- PASS 2: FORMATTING ---")
                format_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": format_prompt}],
                    max_tokens=4000,
                    temperature=0,  # low temperature for consistent formatting
                    top_p=0.95,
                    top_k=20,
                    stream=True,
                )

                # This is the stream we will show to the user
                stream = format_stream
                # --- END: TWO-PASS GENERATION ---

                # Stream to console (unfiltered) and GUI (filtered)
                streamer = TokenStreamer()
                gui_filter = GUITokenFilter()
                
                for chunk in stream:
                    # Console output (unfiltered for debugging)
                    streamer(chunk)
                    
                    # GUI output (filtered)
                    token = chunk["choices"][0]["delta"].get("content","")
                    if token:
                        # Filter token for GUI display
                        clean_token = gui_filter.filter_token(token)
                        if clean_token:
                            self.generated_report_text += clean_token
                            self.after(0, self.update_generation_textbox, clean_token)
                            
                streamer.done()
                
                # Add newlines between reports for different dates
                self.generated_report_text += "\n\n"
                self.after(0, self.update_generation_textbox, "\n\n")

            # Generation complete
            self.after(0, self.generation_finished)
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror(
                "Generation Error",
                f"An error occurred while generating the report: {str(e)}\n\n{traceback.format_exc()}"
            ))
            # Switch back to review view on error
            self.after(0, lambda: (
                self.generation_state.pack_forget(),
                self.review_state.pack(fill="both", expand=True)
            ))
            
        finally:
            # Re-enable original generate button
            self.after(0, lambda: self.generate_btn.configure(
                text="Generate Report",
                state="normal"
            ))
            
    def _extract_clean_summary(self, raw_summary: str) -> str:
        """Extract clean bullet point summaries from raw LLM output, removing thinking tags."""
        # Find and remove the thinking block
        think_end = raw_summary.find('</think>')
        if think_end != -1:
            # Remove everything up to and including </think>
            cleaned = raw_summary[think_end + 8:].strip()
        else:
            # No proper thinking block found, try to work with what we have
            cleaned = raw_summary
        
        # Strip every line just in case
        lines = cleaned.split('\n')
        item_lines = []
        for line in lines:
            stripped = line.strip()
            item_lines.append(stripped)
        return '\n'.join(item_lines)

    def update_generation_textbox(self, token: str):
        """Appends a token to the generation textbox. Must be called from the main thread."""
        self.generation_textbox.insert("end", token)
        self.generation_textbox.see("end")

    def generation_finished(self):
        """Called when the report generation is complete."""
        self.save_btn.configure(state="normal")
        messagebox.showinfo("Generation Complete", "The report has been generated. You can now review it and save the document.")

    def save_generated_report(self):
        """Creates and saves the Word document from the generated text."""
        if not self.generated_report_text.strip():
            messagebox.showwarning("No Content", "There is no generated report to save.")
            return

        try:
            doc = self.create_word_document(self.generated_report_text, self.meeting_dates_for_report)
            self._save_document(doc)
        except Exception as e:
            messagebox.showerror("Save Error", f"An error occurred while creating the document: {str(e)}\n\n{traceback.format_exc()}")

    def create_word_document(self, content, meeting_dates):
        """Create the Word document with proper formatting from raw text."""
        doc = Document()
        
        # Set single spacing for the document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        # Add content
        doc.add_paragraph()  # Blank line
        
        # Updated date
        current_date = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(f"Updated {current_date}")
        
        # Calculate month range
        if meeting_dates:
            try:
                # Date format is 'DD-Mon', e.g., '25-Aug'. Assume current year.
                dates = [datetime.strptime(f"{d}-{datetime.now().year}", "%d-%b-%Y") for d in meeting_dates]
                
                # The list is chronologically sorted, so min is first, max is last.
                min_date = dates[0]
                max_date = dates[-1]

                # Handle year-end crossover (e.g., Dec to Jan)
                if min_date.month > max_date.month:
                    max_date = max_date.replace(year=min_date.year + 1)
                
                if min_date.strftime('%B %Y') == max_date.strftime('%B %Y'):
                    month_range = min_date.strftime("%B %Y")
                elif min_date.year != max_date.year:
                    month_range = f"{min_date.strftime('%B %Y')} - {max_date.strftime('%B %Y')}"
                else:
                    # Same year, different months
                    month_range = f"{min_date.strftime('%B')} - {max_date.strftime('%B, %Y')}"
            except (ValueError, IndexError):
                # Fallback if parsing fails or list is empty
                month_range = datetime.now().strftime("%B %Y")
        else:
            month_range = datetime.now().strftime("%B %Y")
            
        # Title
        title = doc.add_paragraph()
        title.paragraph_format.space_after = Pt(6)
        title_run = title.add_run(f"Major Council Agenda Items, Tentative for {month_range}")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Note
        note_text = ("Note: This is a Tentative Agenda Listing. Dates of items are subject to change "
                    "up to the last minute for a variety of reasons. In addition, this listing does not "
                    "necessarily report all items, just ones that are noteworthy. The City Manager typically "
                    "reviews the tentative agenda items list in more detail with each Councilmember during "
                    "individual meetings.")
        note = doc.add_paragraph()
        note_run = note.add_run(note_text)
        note_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(12)
        
        # Add LLM content by parsing it
        is_first_date = True
        for line in content.split('\n'):
            stripped_line = line.strip()
            if not stripped_line:
                continue

            # Heuristic to detect date line (doesn't start with '-' or a known section header)
            is_date_line = not (
                stripped_line.startswith("- ") or
                stripped_line.startswith("Study Session:") or
                stripped_line.startswith("Closed Session:") or
                stripped_line.startswith("Special Presentations:") or
                stripped_line.startswith("Consent:") or
                stripped_line.startswith("Consideration or Public Hearing:")
            )

            if is_date_line:
                if not is_first_date:
                    doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(18)
                is_first_date = False

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                runner = p.add_run(stripped_line)
                runner.bold = True
                runner.font.size = Pt(14)

            # Check for section headers (Level 1 Bullet)
            elif (stripped_line.startswith("Study Session:") or
                  stripped_line.startswith("Closed Session:") or
                  stripped_line.startswith("Special Presentations:") or
                  stripped_line.startswith("Consent:") or
                  stripped_line.startswith("Consideration or Public Hearing:")):
                p = doc.add_paragraph(stripped_line, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)

            elif stripped_line.startswith("- "):
                # Item under a section (Level 2 Bullet)
                p = doc.add_paragraph(stripped_line[2:].strip(), style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.75)

        # Add horizontal line
        doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(12)
        
        # TBD section
        tbd = doc.add_paragraph("TBD:")
        tbd.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        doc.add_paragraph()
        
        # Significant items section
        last_report_date = (datetime.now() - pd.Timedelta(days=30)).strftime("%B %d, %Y")
        sig_items = doc.add_paragraph(f"Significant Items Completed Since {last_report_date}:")
        sig_items.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        
        return doc
        
    def _save_document(self, doc):
        """Save the document with file dialog"""
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            initialfile=f"Council_Agenda_Summary_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        
        if filename:
            try:
                doc.save(filename)
                messagebox.showinfo(
                    "Success",
                    f"Report generated successfully and saved to:\n{filename}"
                )
            except Exception as e:
                messagebox.showerror(
                    "Save Error",
                    f"Failed to save document: {str(e)}"
                )


def main():
    """Main entry point"""
    app = AgendaSummaryGenerator()
    app.mainloop()


if __name__ == "__main__":
    main()