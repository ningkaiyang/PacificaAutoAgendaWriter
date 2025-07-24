"""
Backend for City of Pacifica – Agenda Summary Generator
Handles:
 • CSV ingestion / validation
 • Local LLM two-pass generation (via llama-cpp-python)
 • Word document creation

The class `AgendaBackend` can be used from any frontend (Tkinter,
Qt, CLI, etc.) without modification.
"""

from __future__ import annotations

import contextlib
import os
import sys
import threading
import time
import traceback
from datetime import datetime
from typing import Callable, List, Sequence

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from llama_cpp import Llama

# Optional – only for simple speed / memory debug
try:
    import resource
except ImportError:  # pragma: no cover  (Windows)
    resource = None  # type: ignore

# --------------------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------------------
def logical_cores() -> int:
    return max(1, os.cpu_count() or 1)


def default_threads() -> int:
    return max(1, logical_cores() // 2)


@contextlib.contextmanager
def suppress_stderr():
    """Silence llama-cpp C-stderr noise while loading the model."""
    with open(os.devnull, "w") as devnull:
        old = sys.stderr
        sys.stderr = devnull
        try:
            yield
        finally:
            sys.stderr = old


class TokenStreamer:
    """Collects streamed tokens, prints all tokens (including thinking tags) for debugging,
    and tracks speed."""
    def __init__(self):
        self._start = time.perf_counter()
        self._tok = 0
        
    def __call__(self, chunk: dict):
        tok = chunk["choices"][0]["delta"].get("content", "")
        if not tok:
            return
        self._tok += 1
        # Print everything to console for debugging
        print(tok, end="", flush=True)
        
    def done(self):
        dt = time.perf_counter() - self._start
        if dt:
            print(f"\nAverage speed: {self._tok/dt:.2f} tok/s")
            print(f"Tokens: {self._tok}")
            print(f"Elapsed Time: {dt:.2f}s")
            # memory usage on macOS is in bytes, so i'll convert to MB
            if resource:
                mem_mb = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss / 1024 / 1024
                print(f"Peak Memory Usage: {mem_mb:.2f} MB")


class GUITokenFilter:
    """
    Strips out <think> … </think> ranges so the GUI only sees "final"
    content. Buffers across token boundaries.
    """
    def __init__(self):
        self._buf = ""
        self._in_think = False

    def filter_token(self, token: str) -> str:
        if not token:
            return ""
        self._buf += token
        out = ""

        while self._buf:
            if not self._in_think:
                start = self._buf.find("<think>")
                if start == -1:
                    out += self._buf
                    self._buf = ""
                else:
                    out += self._buf[:start]
                    self._buf = self._buf[start + len("<think>") :]
                    self._in_think = True
            else:
                end = self._buf.find("</think>")
                if end == -1:
                    # inside a <think>… block, discard until we see the end
                    self._buf = ""
                else:
                    self._buf = self._buf[end + len("</think>") :]
                    self._in_think = False
        return out


# --------------------------------------------------------------------------------------
# Constant definitions (copied from original file)
# --------------------------------------------------------------------------------------
REQUIRED_HEADERS: List[str] = [
    "MEETING DATE", "AGENDA ITEM", "DEPT", "AGENDA SECTION", "ACTION",
    "ATTY REVIEW REQ'D", "FIN REVIEW REQ'D", "DUE TO CAO / FIN", "DRAFT TO CM",
    "MAYOR MTG", "MT FINAL DUE", "AGENDA PUBLISH", "NOTES", "EXCEPTIONS",
    "NOTICING REQ;\nREGULATORY DEADLINES", "CONFLICT OF INTEREST REVIEW NEEDED (Y / N)",
    "HIGHLIGHT IN CWP", "Include in Summary for Mayor"
]

PROMPT_TEMPLATE = """You are an expert city clerk responsible for creating agenda summaries for the City Council. Your task is to take a list of agenda items for a specific meeting date and format them into a clear, concise report.

You have recieved a set of summarized items. You are to categorize them and properly edit them in small ways if necessary (capitalization, merging, etc) and put them together into a resport.
Follow these rules strictly:
1.  Format: The output must be raw text only. Do not use any markdown like '##' or '**'.
2.  Date Header: The report must start with the FULL month name followed by the day number, e.g. "January 1:".  NEVER use numeric-month abbreviations such as "1-Jan".  If there are meeting-level notes, place them in parentheses immediately after the date.
3.  Sections: The report MUST CONTAIN each of these headers ON THEIR OWN LINE, and in the following order:
        "Study Session:"
        "Closed Session:"
        "Special Presentations:"
        "Consent:"
        "Consideration or Public Hearing:"
    If a section has no items, write "TBD" right after the section name. Example: "Study Session: TBD" or "Closed Session: TBD" or "Special Presentations: TBD" or "Consent: TBD" or "Consideration or Public Hearing: TBD".
4.  Item Bullet Points:
    - CRITICAL: Each individual agenda item provided to you MUST be on its own new line in the output.
    - Every item's line must start with a single hyphen and a space: "- ". Do NOT use other bullet point characters like '•' to start off a new line.

Here are some examples of the desired output format:
<examples>
Example 1 (full mix, including a populated Study Session)
September 10:
Study Session:
- Joint Study Session on Revenue Options - ADD DESCRIPTION
Closed Session: TBD
Special Presentations: TBD
Consent:
- Bi-Weekly Disbursements approval
- Approval of Minutes for 1/1/2025 City Council Meeting
Consideration or Public Hearing:
- FY 2025-26 Budget Adoption
- Introduction of Ordinance Changing Council Meeting start-time

Example 2 (showing sections that are entirely TBD and an item needing a description)
July 14:
Study Session: TBD
Closed Session: TBD
Special Presentations:
- Joann Arnos, OSPAC Years of Service
Consent:
- Annual POs/Agreements over $75K PWD-Wastewater
- Sewer Service Charges for FY 2025-26 (last year of approved 5-year schedule)
- Resolution for Park Naming - ADD DESCRIPTION
Consideration or Public Hearing: TBD

NEGATIVE Example (demonstrates what NOT to do — bad bullet characters, bad date format, overly long descriptions, mixed dashes, misplaced headers, messy):
25-Aug:
Closed Session: CLOSED SESSION - TBD • ADD DESCRIPTION - per K.Woodhouse 6/3
Special Presentations:
- City Staff New Hires (Semi-Annual Update) (placeholder) - moved from 6/23 to 8/25 per Y.Carter; HR to provide List of New Hires to K.Woodhouse for review
- Proclamation - Suicide Prevention Month - September 2025 (placeholder) - ADD DESCRIPTION
- Proclamation - National Preparedness Month - September 2025 (placeholder) - ADD DESCRIPTION
Consideration or Public Hearing:
- Housing Element Rezoning EIR Certification + Ordinance Introduction (possibly continued from 8/11) • Final Certification of EIR for Housing Element General Plan Amendments, Rezoning, and Objective Development Standards; Adoption of General Plan amendments; Introduction of Rezoning Ordinance - CAO Review: K.Murphy; To PC 5/19 & 7/7 mtg before going to Council
- Resolution to Amend Council Rules & Code of Ethics to Change City Council Meeting Start Time to 6:00 PM & adopt other outcomes / direction from Council Governance Training (e.g. Vice Mayor nomenclature) • ADD DESCRIPTION - moved from 6/23 per K.Woodhouse; note: Municipal Code refers to Council Rules & Code of Ethics for regular meeting dates / start time and manner of conducting City Council meetings
Public Hearing: Housing Element Rezoning EIR Certification + Ordinance Introduction (possibly continued from 8/11)
Study Session on Revenue Generation (Title TBD from K. Woodhouse) • ADD DESCRIPTION - per K.Woodhouse 6/3
</examples>

Now, using the examples and negative example and given agenda items, generate a report for the following meeting date based on the items provided below. IMPORTANT: List each item under the CORRECT categories, and format the entire agenda CAREFULLY!

Meeting Date: {meeting_date} - IMPORTANT: THIS IS THE ACTUAL METING DATE FOR YOUR REPORT!!! Parse carefully, i.e. "8-Sep" = "September 8"!

Agenda Items:
<items_to_sort>
{items_text}
</items_to_sort>

Report:
"""

# --------------------------------------------------------------------------------------
# Backend main class
# --------------------------------------------------------------------------------------
class AgendaBackend:
    """
    All heavy-lifting lives here.  GUI / CLI frontends interface only via
    the public methods of this class.
    """

    def __init__(self, model_path: str | None = None):
        self.model_path = (
            model_path or "language_models/Qwen3-4B-Q6_K.gguf"
        )
        self.llm_model: Llama | None = None
        self._load_llm_model_async()  # Non-blocking

    # ------------------------------------------------------------------ CSV helpers
    @staticmethod
    def _validate_headers(df: pd.DataFrame) -> bool:
        return all(h in df.columns for h in REQUIRED_HEADERS)

    def process_csv(self, filepath: str) -> tuple[pd.DataFrame, List[pd.Series]]:
        """Read CSV, validate headers, filter rows → return (df, filtered_items)."""
        df = pd.read_csv(filepath)
        if not self._validate_headers(df):
            raise ValueError("Selected CSV is missing required columns.")
        filtered: List[pd.Series] = [
            row for _, row in df.iterrows()
            if str(row.get("Include in Summary for Mayor", "")).upper() == "Y"
        ]
        if not filtered:
            raise RuntimeError("No rows marked 'Y' for inclusion.")
        return df, filtered

    # ------------------------------------------------------------------ LLM loading
    def _load_llm_model_async(self):
        def _loader():
            try:
                if not os.path.exists(self.model_path):
                    print(f"[backend] Model file not found: {self.model_path}")
                    return
                with suppress_stderr():
                    self.llm_model = Llama(
                        model_path=self.model_path,
                        chat_format="chatml",
                        n_ctx=10000,
                        n_threads=default_threads(),
                        verbose=False,
                    )
            except Exception as exc:
                traceback.print_exc()
                print(f"[backend] Failed to load model: {exc}")

        threading.Thread(target=_loader, daemon=True).start()

    # ------------------------------------------------------------------ Public generation API
    def generate_report(
        self,
        selected_rows: Sequence[pd.Series],
        *,
        token_callback: Callable[[str], None] | None = None,
        done_callback: Callable[[str, List[str]], None] | None = None,
        error_callback: Callable[[Exception], None] | None = None,
    ):
        """
        Two-pass streaming generation.
        • token_callback receives GUI-safe text snippets.
        • done_callback(full_report_text, meeting_dates)
        Errors are forwarded to error_callback.
        """

        if not selected_rows:
            raise ValueError("No data passed to backend.generate_report")

        if self.llm_model is None:
            raise RuntimeError("Model not loaded yet – please wait.")

        thread = threading.Thread(
            target=self._run_generation,
            args=(list(selected_rows), token_callback, done_callback, error_callback),
            daemon=True,
        )
        thread.start()

    # ------------------------------------------------------------------ Internal generation logic
    def _run_generation(
        self,
        rows: List[pd.Series],
        token_cb: Callable[[str], None] | None,
        done_cb: Callable[[str, List[str]], None] | None,
        err_cb: Callable[[Exception], None] | None,
    ):
        gui_filter = GUITokenFilter()

        try:
            # Group rows by date, preserving original order
            grouped: dict[str, List[pd.Series]] = {}
            ordered_dates: List[str] = []
            for r in rows:
                date = str(r["MEETING DATE"])
                if date not in grouped:
                    grouped[date] = []
                    ordered_dates.append(date)
                grouped[date].append(r)

            full_output = ""

            for md in ordered_dates:
                items = grouped[md]
                items.sort(key=lambda x: str(x.get("AGENDA SECTION", "")))

                # Build items text for summarisation pass
                items_text = ""
                for it in items:
                    sec = str(it.get("AGENDA SECTION", "N/A")).replace("\n", " ").replace("•", "-").strip()
                    title = str(it.get("AGENDA ITEM", "N/A")).replace("\n", " ").replace("•", "-").strip()
                    notes_val = it.get("NOTES")
                    entry = f"- Item: {title}, Section: \"{sec}\""
                    if pd.notna(notes_val):
                        notes = str(notes_val).replace("\n", " ").replace("•", "-").strip()
                        if notes and notes.lower() != "nan":
                            entry += f", Notes: \"{notes}\""
                    items_text += entry + "\n"

                # ------------ PASS 1 – single-line summaries
                summarization_prompt = f"""You are an expert city clerk. Your task is to summarize each agenda item into ONE short clause.

Rules for summarization:
- Summarize each agenda item in ONE concise single clause as short and clean as possible that clearly signals what the item is. You can omit most parenthesized text from original inputs. Attempt to split or summarize further if it reads like a run-on sentence.
- You should first figure out which category each item belongs in and prepend it to the item: "Study Session:" or "Closed Session:" or "Special Presentations:" or "Consent:" or "Consideration or Public Hearing:". IMPORTANT: ALL considerations OR public hearings go under "Consideration or Public Hearing:".
- You MUST omit unnecessary internal workflow words such as "moved from [dates]", and "per [person]". DO NOT say "moved from 1/1 to 12/31 per Y.Carter" or "per K.Woodhouse" or such.
- If an item INCLUDES the TEXT "placeholder" SPECIFICALLY (NOT the text "TBD", etc), DELETE the entire placeholder and append "(placeholder)" to the end with no other unnecessary placeholder details.
- If an item INCLUDES " ADD DESCRIPTION", DELETE it and append " - ADD DESCRIPTION" to the end, after any potential "(placeholder)".
- Each summary title MUST use Title Case (capitalize all principal words), for example: "Approval of Minutes for 1/1/2025 Meeting".
- If an agenda item includes a workflow date, DELETE it. For logistic dates that belong in the item, keep the date exactly as it appears; do not convert month names or add/remove leading zeros.
- You are also allowed to split long items into seperate items if they would do well as concise seperate items.

Some good examples:
<examples>
Study Session: Study Session on Revenue Generation - ADD DESCRIPTION
Closed Session: TBD - ADD DESCRIPTION
Special Presentations: City Staff New Hires (Semi-Annual Update)
Consent: Annual POs/Agreements over $75K PWD-Wastewater
Consent: Police Militarized Equipment Annual Update - ADD DESCRIPTION (placeholder)
Consent: Sewer service charges for FY2025-26 (last year of approved 5-year schedule)
Consent: Approval of Minutes for 1/1/2024 City Council Meeting
Consideration or Public Hearing: Resolution to Establish Climate Action & Resilience Plan Implementation Committee per CAAP Task Force Charter
Consideration or Public Hearing: Continued Consideration of Climate Action and Resilience Plan Adoption
</examples>

Meeting Date: {date} - IMPORTANT! THIS IS THE ACTUAL MEETING DATE, KEEP TRACK OF IT CAREFULLY! Start off your summarization lines with this meeting date, parsed neatly as <Month Day> like "Meeting Date: January 1" or "Meeting Date: December 31". Parse carefully, i.e. "8-Sep" = "Meeting Date: September 8"!

Agenda Items to Summarize - ONLY SUMMARIZE THESE, DO NOT ADD IN FROM EXAMPLES ACCIDENTALLY:
<summarize_these>
{items_text.strip()}
</summarize_these>

IMPORTANT: Do not explain your reasoning and think too much in circles; output the cleaned lines immediately.
Provide ONLY the proper meeting date format and then the summarized lines, each CAREFULLY capitalized and prepended properly, one per line: /think"""
                
                print("\n--- PASS 1: SUMMARIZATION ---")
                pass1_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": summarization_prompt}],
                    max_tokens=10_000,
                    temperature=0,
                    top_p=0.0,
                    top_k=20,
                    stream=True,
                )
                
                # Collect summarized items from Pass 1
                think_streamer = TokenStreamer()
                raw_summary = ""
                for chunk in pass1_stream:
                    token = chunk["choices"][0]["delta"].get("content", "")
                    raw_summary += token
                    think_streamer(chunk)  # count tokens and print for debug
                think_streamer.done()

                # Clean up summarized_items to remove any incomplete thinking tags
                # and extract only the actual bullet point content
                clean_summary = self._extract_clean_summary(raw_summary)

                # ------------ PASS 2 – final formatting
                # use the main prompt template for the final formatting pass
                # we replace the original items_text with the summarized ones from pass 1
                # and add a /no_think instruction to prevent the model from re-summarizing
                format_prompt = PROMPT_TEMPLATE.format(
                    meeting_date=md, items_text=clean_summary.strip()
                ) + " /no_think"

                print("\n--- PASS 2: FORMATTING ---")
                pass2_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": format_prompt}],
                    max_tokens=10_000,
                    temperature=0,
                    top_p=0.0,
                    top_k=20,
                    stream=True,
                )

                # This is the stream we will show to the user
                # Stream to console (unfiltered) and GUI (filtered)
                streamer = TokenStreamer()
                
                for chunk in pass2_stream:
                    # Console output (unfiltered for debugging)
                    streamer(chunk)
                    
                    # GUI output (filtered)
                    tok = chunk["choices"][0]["delta"].get("content", "")
                    if tok:
                        # Filter token for GUI display
                        cleaned = gui_filter.filter_token(tok)
                        if cleaned and token_cb:
                            token_cb(cleaned)
                            full_output += cleaned
                            
                streamer.done()

                # Separate dates
                if token_cb:
                    token_cb("\n\n")
                full_output += "\n\n"

            streamer.done()
            if done_cb:
                done_cb(full_output, ordered_dates)

        except Exception as exc:  # pragma: no cover
            if err_cb:
                err_cb(exc)
            else:
                traceback.print_exc()



    @staticmethod
    def _extract_clean_summary(raw: str) -> str:
        """Strip everything up to and incl. </think> if present."""
        end = raw.find("</think>")
        cleaned = raw[end + 8 :] if end != -1 else raw
        lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]
        return "\n".join(lines)

    # ------------------------------------------------------------------ Word DOC creation
    @staticmethod
    def create_word_document(content: str, meeting_dates: List[str]) -> Document:
        from datetime import datetime

        doc = Document()
        # Basic style (Calibri 11)
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        doc.add_paragraph()
        doc.add_paragraph(f"Updated {datetime.now():%B %d, %Y}")

        # Month range for title
        month_range = AgendaBackend._month_range(meeting_dates)
        title = doc.add_paragraph()
        title_run = title.add_run(f"Major Council Agenda Items, Tentative for {month_range}")
        title_run.bold = True
        title_run.font.size = Pt(16)

        note_p = doc.add_paragraph()
        note_p.add_run(
            "Note: This is a Tentative Agenda Listing. Dates of items are subject to change "
            "up to the last minute for a variety of reasons. In addition, this listing does not "
            "necessarily report all items, just ones that are noteworthy. The City Manager "
            "typically reviews the tentative agenda items list in more detail with each "
            "Councilmember during individual meetings."
        ).italic = True

        doc.add_paragraph("_" * 78)

        is_first_day = True
        for line in content.splitlines():
            line = line.rstrip()
            if not line:
                continue
            # Date header (heuristic)
            if not line.startswith(("- ", "Study Session:", "Closed Session:",
                                     "Special Presentations:", "Consent:",
                                     "Consideration or Public Hearing:")):
                if not is_first_day:
                    doc.add_paragraph("_" * 78)
                is_first_day = False
                p = doc.add_paragraph()
                p.add_run(line).bold = True
                p.paragraph_format.space_before = Pt(12)
                continue

            # Section header
            if line.startswith(("Study Session:", "Closed Session:",
                                "Special Presentations:", "Consent:",
                                "Consideration or Public Hearing:")):
                p = doc.add_paragraph(line, style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25)
                continue

            # Bullet item
            if line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.75)

        # Closing placeholders
        doc.add_paragraph("_" * 78)
        doc.add_paragraph().add_run("TBD:").bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        last_rep = (datetime.now() - pd.Timedelta(days=60)).strftime("%B %d, %Y")
        doc.add_paragraph().add_run(
            f"Significant Items Completed Since {last_rep}:"
        ).bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")

        return doc

    # Utility for month range human-readable string
    @staticmethod
    def _month_range(dates: List[str]) -> str:
        if not dates:
            return datetime.now().strftime("%B %Y")
        try:
            parsed = [
                datetime.strptime(f"{d}-{datetime.now().year}", "%d-%b-%Y")
                for d in dates
            ]
            parsed.sort()
        except Exception:
            return datetime.now().strftime("%B %Y")

        start, end = parsed[0], parsed[-1]
        if start.strftime("%B %Y") == end.strftime("%B %Y"):
            return start.strftime("%B %Y")
        if start.year != end.year:
            return f"{start:%B %Y} - {end:%B %Y}"
        return f"{start:%B} - {end:%B, %Y}"