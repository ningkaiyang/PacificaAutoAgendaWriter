"""
Backend for City of Pacifica - Agenda Summary Generator
Handles:
 • CSV ingestion / validation
 • Local LLM two-pass generation (via llama-cpp-python)
 • Word document creation

The class `AgendaBackend` can be used from any frontend (Tkinter,
Qt, CLI, etc.) without modification.
"""

from __future__ import annotations

import contextlib
import os
import re
import sys
import threading
import time
import traceback
from datetime import datetime
from typing import Callable, List, Sequence

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from llama_cpp import Llama

# Model to download
MODEL_REPO = "unsloth/Qwen3-4B-GGUF"
MODEL_FILENAME = "Qwen3-4B-Q6_K.gguf"

# resource module is Unix-specific, so we remove it for Windows compatibility.
resource = None

# --------------------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------------------
def logical_cores() -> int:
    return max(1, os.cpu_count() or 1)


def default_threads() -> int:
    return max(1, logical_cores() // 2)


@contextlib.contextmanager
def suppress_stderr():
    """Silence llama-cpp C-stderr noise while loading the model."""
    with open(os.devnull, "w") as devnull:
        old = sys.stderr
        sys.stderr = devnull
        try:
            yield
        finally:
            sys.stderr = old


class TokenStreamer:
    """Collects streamed tokens, prints all tokens (including thinking tags) for debugging,
    and tracks speed."""
    def __init__(self, debug_callback: Callable[[str], None] | None = None):
        self._start = time.perf_counter()
        self._tok = 0
        self.debug_callback = debug_callback

    def __call__(self, chunk: dict):
        tok = chunk["choices"][0]["delta"].get("content", "")
        if not tok:
            return
        self._tok += 1
        # Print everything to console for debugging
        if self.debug_callback:
            self.debug_callback(tok)
        else:
            print(tok, end="", flush=True)

    def done(self):
        dt = time.perf_counter() - self._start
        if dt:
            stats = []
            stats.append(f"\nAverage speed: {self._tok/dt:.2f} tok/s")
            stats.append(f"Tokens: {self._tok}")
            stats.append(f"Elapsed Time: {dt:.2f}s")
            # Peak memory usage reporting is disabled as it's not cross-platform.
            
            stats_str = "\n".join(stats)
            
            if self.debug_callback:
                self.debug_callback(stats_str + "\n")
            else:
                print(stats_str)


class GUITokenFilter:
    """
    Strips out <think> … </think> ranges so the GUI only sees "final"
    content. Buffers across token boundaries.
    """
    def __init__(self):
        self._buf = ""
        self._in_think = False

    def filter_token(self, token: str) -> str:
        if not token:
            return ""
        self._buf += token
        out = ""

        while self._buf:
            if not self._in_think:
                start = self._buf.find("<think>")
                if start == -1:
                    out += self._buf
                    self._buf = ""
                else:
                    out += self._buf[:start]
                    self._buf = self._buf[start + len("<think>") :]
                    self._in_think = True
            else:
                end = self._buf.find("</think>")
                if end == -1:
                    # inside a <think>… block, discard until we see the end
                    self._buf = ""
                else:
                    self._buf = self._buf[end + len("</think>") :]
                    self._in_think = False
        return out


# --------------------------------------------------------------------------------------
# Constant definitions (copied from original file)
# --------------------------------------------------------------------------------------
PROMPT_TEMPLATE_PASS1 = """You are an expert city clerk. Your task is to summarize each agenda item into ONE short clause.

THINK STEP BY STEP, ONCE PER ITEM AND NO MORE. ONCE YOU ARE DONE WITH EVERY ITEM IMMEDIATELY EXIT YOUR THINKING BLOCK AND OUTPUT THE SUMMARIZED LINES.
Rules for summarization:
- Summarize each agenda item in ONE concise single clause as short and clean as possible that clearly signals what the item is. You can omit most parenthesized text from original inputs. Attempt to split or summarize further if it reads like a run-on sentence.
- You should first figure out which category each item belongs in and prepend it to the item: "Study Session:" or "Closed Session:" or "Special Presentations:" or "Consent:" or "Consideration or Public Hearing:". IMPORTANT: ALL considerations OR public hearings go under "Consideration or Public Hearing:".
- You MUST omit unnecessary internal workflow words such as "moved from [dates]", and "per [person]". DO NOT say "moved from 1/1 to 12/31 per Y.Carter" or "per K.Woodhouse" or such.
- If an item INCLUDES the TEXT "placeholder" SPECIFICALLY (NOT the text "TBD", etc), DELETE the entire placeholder and append "(placeholder)" to the end with no other unnecessary placeholder details.
- If an item INCLUDES " ADD DESCRIPTION", DELETE it and append " - ADD DESCRIPTION" to the end, after any potential "(placeholder)".
- Each summary title MUST use Title Case (capitalize all principal words), for example: "Approval of Minutes for 1/1/2025 Meeting".
- If an agenda item includes a workflow date, DELETE it. For logistic dates that belong in the item, keep the date exactly as it appears; do not convert month names or add/remove leading zeros.
- You are also allowed to split long items into seperate items if they would do well as concise seperate items.

Some good examples:
<examples>
Study Session: Study Session on Revenue Generation - ADD DESCRIPTION
Closed Session: TBD - ADD DESCRIPTION
Special Presentations: City Staff New Hires (Semi-Annual Update)
Consent: Annual POs/Agreements over $75K PWD-Wastewater
Consent: Police Militarized Equipment Annual Update - ADD DESCRIPTION (placeholder)
Consent: Sewer service charges for FY2025-26 (last year of approved 5-year schedule)
Consent: Approval of Minutes for 1/1/2024 City Council Meeting
Consideration or Public Hearing: Resolution to Establish Climate Action & Resilience Plan Implementation Committee per CAAP Task Force Charter
Consideration or Public Hearing: Continued Consideration of Climate Action and Resilience Plan Adoption
</examples>

Meeting Date: {md} - IMPORTANT! THIS IS THE ACTUAL MEETING DATE, KEEP TRACK OF IT CAREFULLY! Start off your summarization lines with this meeting date, parsed neatly as <Month Day> like "Meeting Date: January 1" or "Meeting Date: December 31". Parse carefully, i.e. "8-Sep" = "Meeting Date: September 8"!

Agenda Items to Summarize - ONLY SUMMARIZE THESE, DO NOT ADD IN FROM EXAMPLES ACCIDENTALLY:
<summarize_these>
{items_text}
</summarize_these>

Provide ONLY the proper meeting date format and then the summarized lines, each CAREFULLY capitalized and prepended properly, one per line: /think"""

PROMPT_TEMPLATE_PASS2 = """You are an expert city clerk responsible for creating agenda summaries for the City Council. Your task is to take a list of agenda items for a specific meeting date and format them into a clear, concise report.

You have recieved a set of summarized items. You are to categorize them and properly edit them in small ways if necessary (capitalization, merging, etc) and put them together into a resport.
Follow these rules strictly:
1.  Format: The output must be raw text only. Do not use any markdown like '##' or '**'.
2.  Date Header: The report must start with the FULL month name followed by the day number, e.g. "January 1:".  NEVER use numeric-month abbreviations such as "1-Jan".  If there are meeting-level notes, place them in parentheses immediately after the date.
3.  Sections: The report MUST CONTAIN each of these headers ON THEIR OWN LINE, and in the following order:
        "Study Session:"
        "Closed Session:"
        "Special Presentations:"
        "Consent:"
        "Consideration or Public Hearing:"
    If a section has no items, write "TBD" right after the section name. Example: "Study Session: TBD" or "Closed Session: TBD" or "Special Presentations: TBD" or "Consent: TBD" or "Consideration or Public Hearing: TBD".
4.  Item Bullet Points:
    - CRITICAL: Each individual agenda item provided to you MUST be on its own new line in the output.
    - Every item's line must start with a single hyphen and a space: "- ". Do NOT use other bullet point characters like '•' to start off a new line.

Here are some examples of the desired output format:
<examples>
Example 1 (full mix, including a populated Study Session)
September 10:
Study Session:
- Joint Study Session on Revenue Options - ADD DESCRIPTION
Closed Session: TBD
Special Presentations: TBD
Consent:
- Bi-Weekly Disbursements approval
- Approval of Minutes for 1/1/2025 City Council Meeting
Consideration or Public Hearing:
- FY 2025-26 Budget Adoption
- Introduction of Ordinance Changing Council Meeting start-time

Example 2 (showing sections that are entirely TBD and an item needing a description)
July 14:
Study Session: TBD
Closed Session: TBD
Special Presentations:
- Joann Arnos, OSPAC Years of Service
Consent:
- Annual POs/Agreements over $75K PWD-Wastewater
- Sewer Service Charges for FY 2025-26 (last year of approved 5-year schedule)
- Resolution for Park Naming - ADD DESCRIPTION
Consideration or Public Hearing: TBD

NEGATIVE Example (demonstrates what NOT to do — bad bullet characters, bad date format, overly long descriptions, mixed dashes, misplaced headers, messy):
25-Aug:
Closed Session: CLOSED SESSION - TBD • ADD DESCRIPTION - per K.Woodhouse 6/3
Special Presentations:
- City Staff New Hires (Semi-Annual Update) (placeholder) - moved from 6/23 to 8/25 per Y.Carter; HR to provide List of New Hires to K.Woodhouse for review
- Proclamation - Suicide Prevention Month - September 2025 (placeholder) - ADD DESCRIPTION
- Proclamation - National Preparedness Month - September 2025 (placeholder) - ADD DESCRIPTION
Consideration or Public Hearing:
- Housing Element Rezoning EIR Certification + Ordinance Introduction (possibly continued from 8/11) • Final Certification of EIR for Housing Element General Plan Amendments, Rezoning, and Objective Development Standards; Adoption of General Plan amendments; Introduction of Rezoning Ordinance - CAO Review: K.Murphy; To PC 5/19 & 7/7 mtg before going to Council
- Resolution to Amend Council Rules & Code of Ethics to Change City Council Meeting Start Time to 6:00 PM & adopt other outcomes / direction from Council Governance Training (e.g. Vice Mayor nomenclature) • ADD DESCRIPTION - moved from 6/23 per K.Woodhouse; note: Municipal Code refers to Council Rules & Code of Ethics for regular meeting dates / start time and manner of conducting City Council meetings
Public Hearing: Housing Element Rezoning EIR Certification + Ordinance Introduction (possibly continued from 8/11)
Study Session on Revenue Generation (Title TBD from K. Woodhouse) • ADD DESCRIPTION - per K.Woodhouse 6/3
</examples>

Now, using the examples and negative example and given agenda items, generate a report for the following meeting date based on the items provided below. IMPORTANT: List each item under the CORRECT categories, and format the entire agenda CAREFULLY!

Meeting Date: {meeting_date} - IMPORTANT: THIS IS THE ACTUAL METING DATE FOR YOUR REPORT!!! Parse carefully, i.e. "8-Sep" = "September 8"!

Agenda Items:
<items_to_sort>
{items_text}
</items_to_sort>

Report:
"""

# --------------------------------------------------------------------------------------
# Backend main class
# --------------------------------------------------------------------------------------
class AgendaBackend:
    """
    All heavy-lifting lives here.  GUI / CLI frontends interface only via
    the public methods of this class.
    """

    def __init__(self, model_path: str | None = None, user_data_dir: str | None = None):
        self.user_data_dir = user_data_dir
        self.model_path = model_path
        self.llm_model: Llama | None = None
        # if self.model_path and os.path.exists(self.model_path):
        #     self._load_llm_model_async()  # Non-blocking

    # ------------------------------------------------------------------ CSV helpers
    @staticmethod
    def _validate_headers(df: pd.DataFrame, required_headers: List[str]):
        """Check for required headers and raise ValueError if any are missing."""
        missing_headers = [h for h in required_headers if h not in df.columns]
        if missing_headers:
            missing_str = ", ".join(f"'{h}'" for h in missing_headers)
            raise ValueError(f"CSV file is missing required columns: {missing_str}")

    def process_csv(self, filepath: str, csv_headers: dict) -> tuple[pd.DataFrame, List[pd.Series]]:
        """Read CSV, validate headers, filter rows → return (df, all_items)."""
        try:
            # Use the 'python' engine for more robust parsing of potentially
            # irregular CSV files. It can handle rows with differing numbers
            # of columns, which is common in this project's source files.
            df = pd.read_csv(filepath, engine="python")
        except Exception as e:
            # Catch file read errors (e.g., file not found, permission error)
            # and other pandas parsing errors.
            raise ValueError(f"Failed to read or parse CSV file: {e}")

        self._validate_headers(df, list(csv_headers.values()))  # Will raise ValueError if invalid

        # only keep rows where MEETING DATE starts with a digit - actual agenda items
        all_items: List[pd.Series] = []
        for _, row in df.iterrows():
            meeting_date = str(row.get(csv_headers["date"], "")).strip()
            if meeting_date and meeting_date[0].isdigit():
                all_items.append(row)

        if not all_items:
            raise RuntimeError("No valid agenda item rows found in the CSV.")
        return df, all_items

    # ------------------------------------------------------------------ LLM loading
    def _load_llm_model_async(self):
        def _loader():
            try:
                if not self.model_path or not os.path.exists(self.model_path):
                    print(f"[backend] Model file not found: {self.model_path}")
                    return
                with suppress_stderr():
                    self.llm_model = Llama(
                        model_path=self.model_path,
                        chat_format="chatml",
                        n_ctx=10000,
                        n_threads=default_threads(),
                        verbose=False,
                        n_gpu_layers=-1,
                    )
            except Exception as exc:
                traceback.print_exc()
                print(f"[backend] Failed to load model: {exc}")

        threading.Thread(target=_loader, daemon=True).start()

    def download_model(
        self,
        done_callback: Callable[[str], None] | None = None,
        error_callback: Callable[[Exception], None] | None = None,
    ):
        """Downloads model from HuggingFace Hub to user_data_dir."""
        try:
            if not self.user_data_dir:
                raise ValueError("user_data_dir not set in backend")

            models_dir = os.path.join(self.user_data_dir, "models")
            os.makedirs(models_dir, exist_ok=True)

            print(f"[backend] Downloading model to: {models_dir}")

            # Suppress llama.cpp noise during download/load
            with suppress_stderr():
                # Llama.from_pretrained downloads AND loads the model, returning an instance.
                new_llm_instance = Llama.from_pretrained(
                    repo_id=MODEL_REPO,
                    filename=MODEL_FILENAME,
                    local_dir=models_dir,
                    local_dir_use_symlinks=False,  # Use False for better cross-platform/packaging support
                    verbose=False,  # Set to False to avoid duplicate progress info
                    # Add other loading params for consistency
                    chat_format="chatml",
                    n_ctx=10000,
                    n_threads=default_threads(),
                    n_gpu_layers=-1,
                )

            final_model_path = new_llm_instance.model_path
            
            # The model is now loaded, assign it to the backend
            self.llm_model = new_llm_instance
            # Get the actual path string from the instance
            self.model_path = final_model_path

            print(f"[backend] Model downloaded and loaded from: {final_model_path}")
            if done_callback:
                done_callback(final_model_path)  # Pass the string path back to the frontend

        except Exception as e:
            traceback.print_exc()
            if error_callback:
                error_callback(e)

    # ------------------------------------------------------------------ Public generation API
    def generate_report(
        self,
        selected_rows: Sequence[pd.Series],
        *,
        token_callback: Callable[[str], None] | None = None,
        done_callback: Callable[[str, List[str]], None] | None = None,
        error_callback: Callable[[Exception], None] | None = None,
        cancel_event: threading.Event | None = None,
        prompt_template_pass1: str | None = None,
        prompt_template_pass2: str | None = None,
        debug_callback: Callable[[str], None] | None = None,
        ignore_brackets: bool = False,
        csv_headers: dict | None = None,
    ):
        """
        Two-pass streaming generation.
        • token_callback receives GUI-safe text snippets.
        • done_callback(full_report_text, meeting_dates)
        Errors are forwarded to error_callback.
        """

        if not selected_rows:
            raise ValueError("No data passed to backend.generate_report")

        if self.llm_model is None:
            raise RuntimeError("Model not loaded yet - please wait.")

        thread = threading.Thread(
            target=self._run_generation,
            args=(
                list(selected_rows),
                token_callback,
                done_callback,
                error_callback,
                cancel_event,
                prompt_template_pass1,
                prompt_template_pass2,
                debug_callback,
                ignore_brackets,
                csv_headers,
            ),
            daemon=True,
        )
        thread.start()

    # ------------------------------------------------------------------ Internal generation logic
    def _run_generation(
        self,
        rows: List[pd.Series],
        token_cb: Callable[[str], None] | None,
        done_cb: Callable[[str, List[str]], None] | None,
        err_cb: Callable[[Exception], None] | None,
        cancel_event: threading.Event | None = None,
        prompt_template_pass1: str | None = None,
        prompt_template_pass2: str | None = None,
        debug_cb: Callable[[str], None] | None = None,
        ignore_brackets: bool = False,
        csv_headers: dict | None = None,
    ):
        gui_filter = GUITokenFilter()

        # Define default headers here as a fallback if None are passed
        if csv_headers is None:
            csv_headers = {
                "date": "MEETING DATE",
                "section": "AGENDA SECTION",
                "item": "AGENDA ITEM",
                "notes": "NOTES",
                "include": "Include in Summary for Mayor",
            }

        try:
            # Group rows by date, preserving original order
            grouped: dict[str, List[pd.Series]] = {}
            ordered_dates: List[str] = []
            for r in rows:
                date = str(r[csv_headers["date"]])
                if date not in grouped:
                    grouped[date] = []
                    ordered_dates.append(date)
                grouped[date].append(r)

            full_output = ""

            for md in ordered_dates:
                if cancel_event and cancel_event.is_set():
                    if debug_cb:
                        debug_cb("\n[backend] Generation cancelled by user.\n")
                    else:
                        print("\n[backend] Generation cancelled by user.")
                    return
                items = grouped[md]
                items.sort(key=lambda x: str(x.get(csv_headers["section"], "")))

                # Build items text for summarisation pass
                items_text = ""
                for it in items:
                    # Pull section, 'placeholder' if none
                    sec = str(it.get(csv_headers["section"], "N/A")).replace("\n", " ").replace("•", "-").strip()
                    if sec == "nan" or sec == "":
                        sec = "placeholder"
                    # Pull title, 'unnamed item' if none
                    title = str(it.get(csv_headers["item"], "N/A")).replace("\n", " ").replace("•", "-").strip()
                    if title == "nan" or title == "":
                        title = "unnamed item"
                    # Pull notes, does not get added at all to entry if none.
                    notes_val = it.get(csv_headers["notes"])
                    notes = str(notes_val).replace("\n", " ").replace("•", "-").strip()
                    # If ignore brackets, strip from each item only, not across entries.
                    if ignore_brackets:
                        sec = re.sub(r'\[.*?\]', '', sec)
                        title = re.sub(r'\[.*?\]', '', title)
                        notes = re.sub(r'\[.*?\]', '', notes)
                    # Build the entry
                    entry = f"- Item: {title}, Section: \"{sec}\""
                    if (notes.lower() != "nan") and (notes != ""):  # Empty notes not included
                        entry += f", Notes: \"{notes}\""

                    items_text += entry + "\n"

                # ------------ PASS 1 - single-line summaries
                template_pass1 = prompt_template_pass1 or PROMPT_TEMPLATE_PASS1
                summarization_prompt = template_pass1.format(
                    md=md, items_text=items_text.strip()
                )
                
                if debug_cb:
                    debug_cb("\n" + "="*20 + " PASS 1: SUMMARIZATION " + "="*20 + "\n")
                    debug_cb("--- PROMPT INPUT ---\n")
                    debug_cb(summarization_prompt)
                    debug_cb("\n\n--- LLM OUTPUT ---\n")
                else:
                    print("\n--- PASS 1: SUMMARIZATION ---")
                
                pass1_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": summarization_prompt}],
                    max_tokens=10_000,
                    temperature=0,
                    top_p=0.0,
                    top_k=20,
                    stream=True,
                )
                
                # Collect summarized items from Pass 1
                think_streamer = TokenStreamer(debug_callback=debug_cb)
                raw_summary = ""
                for chunk in pass1_stream:
                    if cancel_event and cancel_event.is_set():
                        if debug_cb:
                            debug_cb("\n[backend] Generation cancelled by user.\n")
                        else:
                            print("\n[backend] Generation cancelled by user.")
                        return
                    token = chunk["choices"][0]["delta"].get("content", "")
                    # Stream raw Pass 1 output (including <think> tags) to GUI
                    if token and token_cb:
                        token_cb(token)
                    raw_summary += token
                    think_streamer(chunk)  # count tokens and print for debug
                think_streamer.done()

                # Clean up summarized_items to remove any incomplete thinking tags
                # and extract only the actual bullet point content
                clean_summary = self._extract_clean_summary(raw_summary)

                # ------------ PASS 2 - final formatting
                template_pass2 = prompt_template_pass2 or PROMPT_TEMPLATE_PASS2
                format_prompt = template_pass2.format(
                    meeting_date=md, items_text=clean_summary.strip()
                ) + " /no_think"

                if debug_cb:
                    debug_cb("\n" + "="*20 + " PASS 2: FORMATTING " + "="*20 + "\n")
                    debug_cb("--- PROMPT INPUT ---\n")
                    debug_cb(format_prompt)
                    debug_cb("\n\n--- LLM OUTPUT ---\n")
                else:
                    print("\n--- PASS 2: FORMATTING ---")
                
                pass2_stream = self.llm_model.create_chat_completion(
                    messages=[{"role": "user", "content": format_prompt}],
                    max_tokens=10_000,
                    temperature=0,
                    top_p=0.0,
                    top_k=20,
                    stream=True,
                )

                # This is the stream we will show to the user
                # Stream to console (unfiltered) and GUI (raw)
                streamer = TokenStreamer(debug_callback=debug_cb)
                
                for chunk in pass2_stream:
                    if cancel_event and cancel_event.is_set():
                        if debug_cb:
                            debug_cb("\n[backend] Generation cancelled by user.\n")
                        else:
                            print("\n[backend] Generation cancelled by user.")
                        return
                    # Console output (unfiltered for debugging)
                    streamer(chunk)
                    
                    tok = chunk["choices"][0]["delta"].get("content", "")
                    if tok:
                        # Stream raw token to GUI to show full process
                        if token_cb:
                            token_cb(tok)

                        # Filter token to build the clean report for saving
                        cleaned = gui_filter.filter_token(tok)
                        if cleaned:
                            full_output += cleaned
                            
                streamer.done()

                # Separate dates
                if token_cb:
                    token_cb("\n\n")
                full_output += "\n\n"

            if cancel_event and cancel_event.is_set():
                if debug_cb:
                    debug_cb("\n[backend] Generation cancelled by user.\n")
                else:
                    print("\n[backend] Generation cancelled by user.")
                return
            if done_cb:
                done_cb(full_output, ordered_dates)

        except Exception as exc:  # pragma: no cover
            if err_cb:
                err_cb(exc)
            else:
                traceback.print_exc()

    @staticmethod
    def _extract_clean_summary(raw: str) -> str:
        """Strip everything up to and incl. </think> if present."""
        end = raw.find("</think>")
        cleaned = raw[end + 8 :] if end != -1 else raw
        lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]
        return "\n".join(lines)

    # ------------------------------------------------------------------ Word DOC creation
    @staticmethod
    def create_word_document(content, meeting_dates):
        """Create the Word document with proper formatting from raw text."""
        doc = Document()
        
        # Set single spacing for the document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        # Add content
        doc.add_paragraph()  # Blank line
        
        # Updated date
        current_date = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(f"Updated {current_date}")
        
        # Calculate month range
        if meeting_dates:
            try:
                # Date format is 'DD-Mon', e.g., '25-Aug'. Assume current year.
                dates = [datetime.strptime(f"{d}-{datetime.now().year}", "%d-%b-%Y") for d in meeting_dates]
                
                # The list is chronologically sorted, so min is first, max is last.
                min_date = dates[0]
                max_date = dates[-1]

                # Handle year-end crossover (e.g., Dec to Jan)
                if min_date.month > max_date.month:
                    max_date = max_date.replace(year=min_date.year + 1)
                
                if min_date.strftime('%B %Y') == max_date.strftime('%B %Y'):
                    month_range = min_date.strftime("%B %Y")
                elif min_date.year != max_date.year:
                    month_range = f"{min_date.strftime('%B %Y')} - {max_date.strftime('%B %Y')}"
                else:
                    # Same year, different months
                    month_range = f"{min_date.strftime('%B')} - {max_date.strftime('%B, %Y')}"
            except (ValueError, IndexError):
                # Fallback if parsing fails or list is empty
                month_range = datetime.now().strftime("%B %Y")
        else:
            month_range = datetime.now().strftime("%B %Y")
            
        # Title
        title = doc.add_paragraph()
        title.paragraph_format.space_after = Pt(6)
        title_run = title.add_run(f"Major Council Agenda Items, Tentative for {month_range}")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Note
        note_text = ("Note: This is a Tentative Agenda Listing. Dates of items are subject to change "
                    "up to the last minute for a variety of reasons. In addition, this listing does not "
                    "necessarily report all items, just ones that are noteworthy. The City Manager typically "
                    "reviews the tentative agenda items list in more detail with each Councilmember during "
                    "individual meetings.")
        note = doc.add_paragraph()
        note_run = note.add_run(note_text)
        note_run.italic = True
        
        # Add horizontal line
        doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(12)
        
        # Add LLM content by parsing it
        is_first_date = True
        for line in content.split('\n'):
            stripped_line = line.strip()
            if not stripped_line:
                continue

            # Heuristic to detect date line (doesn't start with '-' or a known section header)
            is_date_line = not (
                stripped_line.startswith("- ") or
                stripped_line.startswith("Study Session:") or
                stripped_line.startswith("Closed Session:") or
                stripped_line.startswith("Special Presentations:") or
                stripped_line.startswith("Consent:") or
                stripped_line.startswith("Consideration or Public Hearing:")
            )

            if is_date_line:
                if not is_first_date:
                    doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(18)
                is_first_date = False

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                runner = p.add_run(stripped_line)
                runner.bold = True
                runner.font.size = Pt(14)

            # Check for section headers (Level 1 Bullet)
            elif (stripped_line.startswith("Study Session:") or
                  stripped_line.startswith("Closed Session:") or
                  stripped_line.startswith("Special Presentations:") or
                  stripped_line.startswith("Consent:") or
                  stripped_line.startswith("Consideration or Public Hearing:")):
                p = doc.add_paragraph(stripped_line, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)

            elif stripped_line.startswith("- "):
                # Item under a section (Level 2 Bullet)
                p = doc.add_paragraph(stripped_line[2:].strip(), style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.75)

        # Add horizontal line
        doc.add_paragraph("_" * 78).paragraph_format.space_before = Pt(12)
        
        # TBD section
        tbd = doc.add_paragraph("TBD:")
        tbd.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        doc.add_paragraph()
        
        # Significant items section
        last_report_date = (datetime.now() - pd.Timedelta(days=60)).strftime("%B %d, %Y")
        sig_items = doc.add_paragraph(f"Significant Items Completed Since {last_report_date}:")
        sig_items.runs[0].bold = True
        doc.add_paragraph("[Placeholder for user to manually enter items.]")
        
        return doc